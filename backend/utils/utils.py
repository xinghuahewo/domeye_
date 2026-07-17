import sys
import os
import datetime
import networkx as nx
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import pandas as pd
import ipaddress
from netaddr import *
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor 
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


from utils.get_as_info import get_as_name, get_as_country_cn, get_as_type, get_as_org_name, get_global_rank, get_country_rank
from utils.get_domain_info import get_domain_ip, get_domain_industry, get_domain_prefix

plt.rcParams['font.sans-serif'] = ['WenQuanYi Zen Hei']  # 使用 SimHei 字体
plt.rcParams['axes.unicode_minus'] = False  # 解决坐标轴负号显示问题


def draw_leak_graph(as_path_str, leak_by_as, leak_to_as, 
                               output_img="leak_graph.png"):
    """
        生成路由泄露示意图
    """
    as_path_list = as_path_str.split(" ")
    as_path_list.reverse()
    # 去除重复
    as_set = set()
    as_path = [x for x in as_path_list if not (x in as_set or as_set.add(x))]


    G = nx.DiGraph()
    
    origin_as = as_path[0]

    # 添加节点和边
    for i in range(len(as_path)-1):
        G.add_edge(f"AS{as_path[i]}", f"AS{as_path[i+1]}")

    # 两边平，中间谷底
    leak_index = as_path.index(leak_by_as)
    pos = {}
    for i, asn in enumerate(as_path):
        x = i * 2  # 增加x轴间距
        y = -2 if i == leak_index else 0  # 泄露点下沉
        pos[f"AS{asn}"] = (x, y)

    # 设置节点属性（使用中文标签）
    node_colors = []
    
    for asn in as_path:
        if asn == as_path[0]:  # 初始节点
            node_colors.append("limegreen")
        elif asn == leak_by_as:  # 泄露方
            node_colors.append("red")
        elif asn == leak_to_as:  # 传播方
            node_colors.append("gold")
        else:  # 其他节点
            node_colors.append("lightgray")

    # 创建图形
    plt.figure(figsize=(15, 6))
    ax = plt.gca()
    
    # 绘制边（带正确方向的箭头）
    for i in range(len(as_path)-1):
        start_node = f"AS{as_path[i]}"
        end_node = f"AS{as_path[i+1]}"
        
        # 计算节点边缘的连接点
        start_pos = pos[start_node]
        end_pos = pos[end_node]
        
        # 向量计算（从起点指向终点）
        dx = end_pos[0] - start_pos[0]
        dy = end_pos[1] - start_pos[1]
        dist = (dx ** 2 + dy ** 2) ** 0.5
        
        # 调整终点到节点边缘（减去节点半径）
        end_adj = (
            end_pos[0] - (dx/dist)*0.4,  # 0.4是节点半径的视觉等效值
            end_pos[1] - (dy/dist)*0.4
        )
        
        # 绘制箭头
        ax.annotate("",
            xy=end_adj, 
            xytext=start_pos,
            arrowprops=dict(
                arrowstyle="->",
                color="#4a8bc9",
                linewidth=2,
                shrinkA=0,  # 起点不收缩
                shrinkB=0,  # 终点不收缩
                # connectionstyle="arc3,rad=0.2"
            ))
        
    # 绘制节点
    for i, asn in enumerate(as_path):
        node = f"AS{asn}"
        nx.draw_networkx_nodes(
            G, pos,
            nodelist=[node],
            node_color=[node_colors[i]],
            node_size=2500,
            edgecolors="black",
            linewidths=2,
            ax=ax
        )
        
        if asn == origin_as:
            node_info = f"受害方：AS{asn}"
        elif asn == leak_by_as:
            node_info = f"泄漏方：AS{asn}"
        elif asn == leak_to_as:
            node_info = f"传播方：AS{asn}"
        else:
            node_info = f"AS{asn}"

        ax.text(pos[node][0], pos[node][1]+0.5, node_info, 
                ha='center', fontsize=12, fontweight='bold')

    # # 添加图例（使用中文标签）
    # legend_elements = [
    #     patches.Patch(facecolor='limegreen', label='origin'),
    #     patches.Patch(facecolor='red', label='leak_by'),
    #     patches.Patch(facecolor='gold', label='leak_to'),
    #     patches.Patch(facecolor='lightgray', label='other')
    # ]
    # ax.legend(handles=legend_elements, loc='upper right')

    plt.xlim(-2, (len(as_path)-1)*2+2)
    plt.ylim(-3, 3)
    plt.axis('off')
    # plt.title("BGPRouteLeak Propagation Diagram", fontsize=16, pad=20)
    
    # 保存图片
    plt.savefig(output_img, dpi=300, bbox_inches="tight", facecolor='white')
    plt.close()
    return output_img

def text(doc, para):
    """
    设置正文格式（仿宋_GB2312 小四）
    """
    paragrapha = doc.add_paragraph(para)
    # 小四 24
    paragrapha.style.font.size = Pt(12)
    paragrapha.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符（14磅×2）
    paragrapha.style.font.name = 'Times New Roman'  # 西文字体
    paragrapha.style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')  # 中文字体
    paragrapha.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    paragrapha.paragraph_format.line_spacing = 1.15  # 设置行间距为1.15
    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 两端对齐

    # 设置段落间距（段前和段后）
    # style_content.paragraph_format.space_before = Pt(6)  # 段前间距 6 磅
    paragrapha.paragraph_format.space_after = Pt(6)   # 段后间距 6 磅

def title(doc, title):
    """
    设置正文格式（仿宋_GB2312 四号）
    """
    paragrapha = doc.add_paragraph()
    # 小四 24
    paragrapha.style.font.size = Pt(28)  # 四号字对应14磅
    # paragrapha.paragraph_format.first_line_indent = Pt(12)  # 首行缩进2字符（14磅×2）
    paragrapha.style.font.name = 'Times New Roman'  # 西文字体
    paragrapha.style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')  # 中文字体
    paragrapha.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    paragrapha.paragraph_format.line_spacing = 1.15  # 设置行间距为1.15
    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 两端对齐

def appendix(doc, para):
    """
    设置正文格式（仿宋_GB2312 四号）
    """
    paragrapha = doc.add_paragraph(para)
    # 小四 24
    paragrapha.style.font.size = Pt(14)  # 四号字对应14磅
    # paragrapha.paragraph_format.first_line_indent = Pt(12)  # 首行缩进2字符（14磅×2）
    paragrapha.style.font.name = 'Times New Roman'  # 西文字体
    paragrapha.style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')  # 中文字体
    paragrapha.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    paragrapha.paragraph_format.line_spacing = 1.15  # 设置行间距为1.15
    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 两端对齐

def heading_1(doc, heading):
    """
    处置模板的一级标题（主标题）
    """
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run(heading)
    run_h1.font.name = '仿宋'
    r = run_h1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run_h1.font.bold = True
    run_h1.font.size = Pt(22)
    h1.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def heading_2(doc, heading):
    """
    处置模板的二级标题
    """
    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run(heading)
    run_h2.font.name = '宋体'
    r = run_h2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_h2.font.bold = True
    run_h2.font.size = Pt(15)
    h2.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

def heading_3(doc, heading):
    """
    处置模板的三级标题
    """
    h3 = doc.add_heading(level=3)
    run_h3 = h3.add_run(heading)
    run_h3.font.name = '宋体'
    r = run_h3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_h3.font.bold = True
    run_h3.font.size = Pt(14)
    h3.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

def subtitle(doc):
    p0 = doc.add_paragraph()
    create_time = str(datetime.datetime.now().strftime('%Y年%m月%d日'))
    run_p0 = p0.add_run('国家互联网应急中心（CNCERT）              {}'.format(create_time))
    run_p0.font.name = 'Times New Roman'
    r = run_p0.element.rPr.rFonts.set(qn('w:eastAsia'), '方正黑体简体')
    run_p0.font.size = Pt(14)
    # run_p0.underline = True
    p0.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def table_1(doc, domain, domain_num, prefix, is_auth, domain_info):
    if domain_num > 5:
        if is_auth:
            p1 = """
        前缀包含的部分权威IP及权威域名服务：
                 """
            text(doc, p1)
        else:
            p1 = """
        前缀包含的部分IP及域名服务：
                 """
            text(doc, p1)
        table1 = doc.add_table(rows=6, cols=6, style="Table Grid")
        table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
        table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "IP", "域名", "服务名称", "分类", "归属路由前缀"
        table1.style.font.size = Pt(12)  # 字体大小15磅
        for i in range(1, 6):
            table1.cell(i, 0).text = str(i)
            table1.cell(i, 2).text = domain[i-1]
            if domain_info.get(domain[i-1]):
                domain_data = domain_info[domain[i-1]]
                table1.cell(i, 1).text = check_ip_in_subnet(prefix=prefix, 
                                                            is_auth=is_auth,
                                                            domain_data=domain_data)
                table1.cell(i, 3).text = domain_data['title'] 
                table1.cell(i, 4).text = domain_data['industry']
            table1.cell(i, 5).text = prefix
        table1.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 字体颜色
        table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
    else:
        if is_auth:
            p1 = """
        前缀包含的权威IP及权威域名服务：
                 """
            text(doc, p1)
        else:
            p1 = """
        前缀包含的IP及域名服务：
                 """
            text(doc, p1)
        rows = domain_num + 1
        table1 = doc.add_table(rows=rows, cols=6, style="Table Grid")
        table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
        table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "IP", "域名", "服务名称", "分类", "归属路由前缀"
        table1.style.font.size = Pt(12)  # 字体大小15磅
        for i in range(1, rows):
            table1.cell(i, 0).text = str(i)
            table1.cell(i, 2).text = domain[i-1]
            if domain_info.get(domain[i-1]):
                domain_data = domain_info[domain[i-1]]
                table1.cell(i, 1).text = check_ip_in_subnet(prefix=prefix, 
                                                            is_auth=is_auth,
                                                            domain_data=domain_data)         
                table1.cell(i, 3).text = domain_data['title'] 
                table1.cell(i, 4).text = domain_data['industry']
            table1.cell(i, 5).text = prefix
        table1.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 字体颜色
        table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐

def check_ip_in_subnet(prefix, is_auth, domain_data):
    prefix_obj = IPNetwork(prefix)
    domain_ip = eval(domain_data['auth_ip']) if is_auth else eval(domain_data['ip'])
    ip_list = list()
    for ip in domain_ip:
        try:
            ip_obj = IPAddress(ip)
            if ip_obj in prefix_obj:
                ip_list.append(ip)
        except:
            pass
    return str(ip_list) if ip_list else ''


def generateTableInfo(doc, para1, level, style_title, style_content):
    """生成第一行  一行三列的表格
    
        eg. 事件编号	CNCERT-BGPMon-Hijack-103.158.37.0/24-2-20250402	高危事件
    """
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    table.cell(0, 0).text = "事件编号"
    table.cell(0, 0).width = Cm(2.5)
    table.cell(0, 0).paragraphs[0].style = style_title

    table.cell(0, 1).text = para1
    table.cell(0, 1).width = Cm(11)
    table.cell(0, 1).paragraphs[0].style = style_content

    table.cell(0, 2).text = level
    table.cell(0, 2).width = Cm(2.5)
    table.cell(0, 2).paragraphs[0].style = style_title


def setTableDetailFormat(table, columns, style_title, style_content, type):
    """设置详细内容的格式与样式"""
    # 设置列宽（必须设置整列宽度才有效）
    for row in table.rows:
        for cell in row.cells:
            item = cell.text.strip()
            if item in columns:
                ## 第一列宽度
                cell.width = Cm(2.5)
                cell.paragraphs[0].style = style_title
            else:
                ## 第二列宽度
                cell.width = Cm(13.5)
                # 画图不需要设置格式
                if item == "路由泄露示意图":
                    continue
                cell.paragraphs[0].style = style_content

    # 设置行高
    for row_idx, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 2:
            continue
            
        item = cells[0].text.strip()
        
        # 设置行高（需要同时设置高度和高度规则）
        if item == "监测情况":
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(3)
        elif item == "影响范围":
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            if type == "hijack":
                row.height = Cm(4)
            elif type == "sub_hijack":
                row.height = Cm(4)
            elif type == "leak":
                row.height = Cm(4)
            elif type == "outage":
                row.height = Cm(7)
        elif item in ["被劫持方情况介绍", "劫持方情况介绍", "路由中断方介绍"]:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(3)
        elif item == "路由泄露涉事方介绍":
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(6.5)
        elif item == "应急处置建议":
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(1.5)
    
    # 设置单元格字体格式
    for row in table.rows:
        for cell in row.cells:
            item = cell.text.strip()
            if item in columns:
                cell.paragraphs[0].style = style_title
            else:
                if item == "路由泄露示意图":
                    continue
                cell.paragraphs[0].style = style_content



    # for row in table.rows:
        # for cell in row.cells:
        #     item = cell.text.strip()

        #     if item in columns:
        #         ## 第一列宽度
        #         cell.width = Cm(2.5)
        #         cell.paragraphs[0].style = style_title
        #     else:
        #         ## 第二列宽度
        #         cell.width = Cm(13.5)
        #         # 画图不需要设置格式
        #         if item == "路由泄露示意图":
        #             continue
        #         cell.paragraphs[0].style = style_content

def generateTbableASInfo(doc, as_info, asn_set, style_content):
    """生成AS信息表格
        asn, as_name, as_country_cn, type, org_name_cn, global_rank, country_rank
    """
    # 在第二页写
    doc.add_page_break()
    para2 = f"附：涉事AS信息"
    appendix(doc, para2)

    table = doc.add_table(rows=len(asn_set)+1, cols=7, style="Table Grid")
    table.autofit = True

    # 设置列宽（整列宽度，而不只是表头）
    column_widths = [
        Cm(2),  # ASN
        Cm(4),  # AS名称
        Cm(2),  # 所属国家
        Cm(2),  # AS类型
        Cm(2),  # 所属组织
        Cm(2),  # 全球排名
        Cm(2)   # 国内排名
    ]
    
    for col_idx, width in enumerate(column_widths):
        for row in table.rows:
            row.cells[col_idx].width = width

    # 设置表头
    headers = ["ASN", "AS名称", "所属国家", "AS类型", "所属组织", "全球排名", "国内排名"]
    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header


    # 设置表格内容
    line_count = 1
    for asn in asn_set:
        asn_str = str(asn)
        as_name = get_as_name(as_info, asn)
        as_country = get_as_country_cn(as_info, asn)
        as_type = get_as_type(as_info, asn)
        as_org = get_as_org_name(as_info, asn)
        as_global_rank = get_global_rank(as_info, asn)
        as_country_rank = get_country_rank(as_info, asn)

        # 添加表格内容
        table.rows[line_count].cells[0].text = asn_str
        table.rows[line_count].cells[1].text = as_name if as_name else ""
        table.rows[line_count].cells[2].text = as_country if as_country else ""
        table.rows[line_count].cells[3].text = as_type if as_type else ""
        table.rows[line_count].cells[4].text = as_org if as_org else ""
        table.rows[line_count].cells[5].text = as_global_rank if as_global_rank else ""
        table.rows[line_count].cells[6].text = as_country_rank if as_country_rank else ""

        line_count += 1
        # 设置单元格格式
    # 修改表格样式
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].style = style_content


def generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, style_content):
    """涉事前缀包含的域名信息"""
    if domain_num > 0 or domain_auth_num > 0:
    
        # # 在第二页写
        # doc.add_page_break()
        para2 = f"附：被劫持前缀包含的域名信息"
        appendix(doc, para2)
        table3 = doc.add_table(rows=domain_auth_num+domain_num+1, cols=5, style="Table Grid")
        table3.autofit = True

        # 设置列宽（整列宽度）
        column_widths = [
            Cm(2),  # 序号
            Cm(4),  # IP地址
            Cm(3),  # 承载信息
            Cm(3),  # 类型
            Cm(4)   # 前缀
        ]
        
        for col_idx, width in enumerate(column_widths):
            for row in table3.rows:
                row.cells[col_idx].width = width

        # 设置表头
        headers = ["序号", "IP地址", "承载信息", "类型", "前缀"]
        for col_idx, header in enumerate(headers):
            table3.cell(0, col_idx).text = header

        line_count = 1

        for item in domain:
            
            item_ip = get_domain_ip(domain_info, item)
            item_prefix = get_domain_prefix(domain_info, item)
            item_industry = get_domain_industry(domain_info, item)
            
            table3.rows[line_count].cells[0].text = str(line_count)
            table3.rows[line_count].cells[1].text = item_ip
            table3.rows[line_count].cells[2].text = item
            table3.rows[line_count].cells[3].text = "域名"
            table3.rows[line_count].cells[4].text = item_prefix
            line_count += 1
        
        for item in domain_auth:
            item_ip = get_domain_ip(domain_info, item)
            item_prefix = get_domain_prefix(domain_info, item)
            item_industry = get_domain_industry(domain_info, item)

            table3.rows[line_count].cells[0].text = str(line_count)
            table3.rows[line_count].cells[1].text = item_ip
            table3.rows[line_count].cells[2].text = item
            table3.rows[line_count].cells[3].text = "权威解析服务"
            table3.rows[line_count].cells[4].text = item_prefix
            line_count += 1

        # 修改表格样式
        for row in table3.rows:
            for cell in row.cells:
                cell.paragraphs[0].style = style_content
