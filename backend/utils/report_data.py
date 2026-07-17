# -*- coding: utf-8 -*-
"""
报告数据聚合服务
提供各类事件报告所需的数据聚合功能
"""

import datetime
from datetime import timedelta
import traceback
import os
import pytz
import json
import ipaddress
from functools import lru_cache
from collections import defaultdict
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib import pyplot as plt
except Exception:
    plt = None

from config.config import PFX2AS_DICT_FILE, BIG_COUNTRY, FEATURE_OTHER_TABLE, SOURCE
from database.prefix_outage import get_outage_prefixes_by_asn_at, select_prefix_outage_by_interval
from database.as_outage import select_as_outage_by_interval
from database.feature_asn import get_as_baseline

from database.country_outage import (
    get_country_outage_de,
    get_country_as_outage_details,
    get_country_prefix_outage_details,
    get_country_allocated_as_count
)
from database.feature_country import (
    get_country_baseline,
    get_country_current,
    get_country_time_series
)

# 日志开关：仅用于定位 v4_total/v6_total 为 0 的原因
REPORT_DEBUG = str(os.environ.get("DOMEYE_REPORT_DEBUG", "")).lower() in ("1", "true", "yes", "y", "on")


def _debug_log(msg: str):
    if REPORT_DEBUG:
        print(f"[report_data][debug] {msg}")


# IPv6 /48 地址块包含的地址数（128-48=80）
V6_ADDR_PER_48 = 1 << 80

# 报告数据查询窗口：默认将国家事件窗口向后扩展 1 天
REPORT_LOOKAHEAD = timedelta(days=1)

# 表格字号：四号（14pt）
TABLE_FONT_SIZE = Pt(14)

# 国家时区映射
COUNTRY_TIMEZONE = {
    'IR': 'Asia/Tehran',
    'CN': 'Asia/Shanghai',
    'US': 'America/New_York',
    'RU': 'Europe/Moscow',
    'JP': 'Asia/Tokyo',
    'KR': 'Asia/Seoul',
    'DE': 'Europe/Berlin',
    'GB': 'Europe/London',
    'FR': 'Europe/Paris',
    'AU': 'Australia/Sydney',
    'IN': 'Asia/Kolkata',
    'BR': 'America/Sao_Paulo',
    'CA': 'America/Toronto',
    'MX': 'America/Mexico_City',
    'AF': 'Asia/Kabul',
    'PK': 'Asia/Karachi',
    'SA': 'Asia/Riyadh',
    'AE': 'Asia/Dubai',
    'EG': 'Africa/Cairo',
    'ZA': 'Africa/Johannesburg',
}

# 国家中文名到代码的映射
COUNTRY_NAME_TO_CODE = {
    '伊朗': 'IR',
    '中国': 'CN',
    '美国': 'US',
    '俄罗斯': 'RU',
    '日本': 'JP',
    '韩国': 'KR',
    '德国': 'DE',
    '英国': 'GB',
    '法国': 'FR',
    '澳大利亚': 'AU',
    '印度': 'IN',
    '巴西': 'BR',
    '加拿大': 'CA',
    '墨西哥': 'MX',
    '阿富汗': 'AF',
    '巴基斯坦': 'PK',
    '沙特阿拉伯': 'SA',
    '阿联酋': 'AE',
    '埃及': 'EG',
    '南非': 'ZA',
}


def _set_run_fonts(run, cn_font: str = "宋体", en_font: str = "Times New Roman", size: Optional[Pt] = None):
    if run is None:
        return
    if size is not None:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:cs"), en_font)


def _set_paragraph_first_line_indent_chars(paragraph, chars: int = 2):
    """
    Word 的“首行缩进 2 字符”是相对字符的缩进，优先用 firstLineChars 实现。
    chars=2 => w:firstLineChars="200"
    """
    if paragraph is None:
        return
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(int(chars) * 100))
    # 避免 firstLine 与 firstLineChars 冲突
    if ind.get(qn("w:firstLine")) is not None:
        ind.attrib.pop(qn("w:firstLine"), None)


def _apply_doc_base_styles(doc: Document, body_size_pt: int = 18):
    """
    统一：中文宋体，英文 Times New Roman，正文小二（18pt）。
    """
    if doc is None:
        return
    body_size = Pt(body_size_pt)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.size = body_size if style_name == "Normal" else style.font.size
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")

        if style_name == "Normal":
            # 正文首行缩进仅在段落级别设置
            ppr = style._element.get_or_add_pPr()
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                ind.attrib.pop(qn("w:firstLineChars"), None)
                ind.attrib.pop(qn("w:firstLine"), None)


def _draw_line_chart(
    *,
    xs,
    ys,
    title: str,
    ylabel: str,
    out_path: str,
):
    if plt is None:
        return None
    if not xs or not ys:
        return None
    fig = plt.figure(figsize=(7.2, 3.6), dpi=160)
    ax = fig.add_subplot(111)
    ax.plot(xs, ys, linewidth=2)
    ax.set_title(title, fontsize=12)
    ax.set_ylabel(ylabel, fontsize=10)
    ax.grid(True, linestyle="--", alpha=0.4)
    fig.autofmt_xdate(rotation=30, ha="right")
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)
    return out_path


def _build_visible_addr_figures(time_series: list[dict], out_dir: str, file_prefix: str):
    """
    生成：
    - 图1：IPv6 /48 前缀数量变化
    - 图2：IPv4 可见地址数量变化
    """
    if not time_series:
        return None, None
    os.makedirs(out_dir, exist_ok=True)

    xs = []
    v6 = []
    v4 = []
    for row in time_series:
        t = row.get("t")
        try:
            dt = datetime.datetime.strptime(str(t), "%Y-%m-%d %H:%M:%S")
        except Exception:
            try:
                dt = datetime.datetime.fromisoformat(str(t))
            except Exception:
                continue
        xs.append(dt)
        v6.append(int(row.get("v6Prefix_num", 0) or 0))
        v4.append(int(row.get("v4IP_num", 0) or 0))

    fig1 = os.path.join(out_dir, f"{file_prefix}_fig1_v6prefix.png")
    fig2 = os.path.join(out_dir, f"{file_prefix}_fig2_v4addr.png")

    fig1_path = _draw_line_chart(
        xs=xs,
        ys=v6,
        title="IPv6 /48 Prefix Count (Visible) Trend",
        ylabel="Prefixes (/48)",
        out_path=fig1,
    )
    fig2_path = _draw_line_chart(
        xs=xs,
        ys=v4,
        title="IPv4 Address Count (Visible) Trend",
        ylabel="Addresses",
        out_path=fig2,
    )
    return fig1_path, fig2_path


def convert_to_local_time(beijing_time, country_code):
    """
    将北京时间转换为指定国家的当地时间
    
    :param beijing_time: 北京时间 (datetime)
    :param country_code: 国家两字母代码
    :return: (local_time_str, timezone_name)
    """
    if isinstance(beijing_time, str):
        beijing_time = datetime.datetime.strptime(beijing_time, "%Y-%m-%d %H:%M:%S")
    
    beijing_tz = pytz.timezone('Asia/Shanghai')
    local_tz_name = COUNTRY_TIMEZONE.get(country_code, 'UTC')
    local_tz = pytz.timezone(local_tz_name)
    
    # 设置北京时间时区
    beijing_dt = beijing_tz.localize(beijing_time)
    # 转换到当地时区
    local_dt = beijing_dt.astimezone(local_tz)
    
    local_time_str = local_dt.strftime("%Y年%m月%d日 %H时%M分")
    return local_time_str, local_tz_name


def format_beijing_time(dt):
    """格式化北京时间为中文格式"""
    if isinstance(dt, str):
        dt = datetime.datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
    return dt.strftime("%Y年%m月%d日 %H时%M分")


def calculate_drop_percent(baseline, current):
    """计算降幅百分比"""
    if baseline == 0:
        return 0
    return round((baseline - current) / baseline * 100, 2)


def get_as_info_from_dict(as_info, asn):
    """从as_info字典获取AS信息"""
    asn_str = str(asn)
    if asn_str in as_info:
        info = as_info[asn_str]
        return {
            'asn': asn_str,
            'as_name': info.get('as_name', ''),
            'org_name': info.get('org_name_cn', '') or info.get('org_name', ''),
            'country': info.get('as_country_cn', ''),
            'v4Prefixes_num': info.get('v4Prefixes_num', 0),
            'v6Prefixes_num': info.get('v6Prefixes_num', 0)
        }
    return {
        'asn': asn_str,
        'as_name': '',
        'org_name': '',
        'country': '',
        'v4Prefixes_num': 0,
        'v6Prefixes_num': 0
    }


def _split_prefixes_by_version(prefixes: set):
    """将前缀集合拆分为 IPv4/IPv6 两个集合（忽略无效前缀）。"""
    v4, v6 = set(), set()
    for p in prefixes:
        if not p:
            continue
        try:
            net = ipaddress.ip_network(str(p).strip(), strict=False)
        except Exception:
            continue
        if net.version == 4:
            v4.add(str(net))
        elif net.version == 6:
            v6.add(str(net))
    return v4, v6


def _prefix_addr_num(prefix: str):
    """
    计算单条前缀对应的地址数（返回字符串，避免前端精度丢失）。

    Returns:
        tuple[int|None, str]: (ip_version, addr_num_str)
    """
    if not prefix:
        return None, "0"
    try:
        net = ipaddress.ip_network(str(prefix).strip(), strict=False)
        return net.version, str(int(net.num_addresses))
    except Exception:
        return None, "0"


def cal_v4_addr_from_prefixes(prefixes: set) -> int:
    """
    IPv4 前缀集合折算为地址数：
    - 将所有 IPv4 前缀按 /24 粒度去重，返回 /24 数量 * 256
    注意：与 core/BGPFeature.calculate_c_segments_count 口径一致：更细粒度（/25+/32）
    也会折算到其所在的 /24。
    """
    v4_24_set = set()
    for prefix in prefixes:
        try:
            net = ipaddress.ip_network(str(prefix).strip(), strict=False)
        except Exception:
            continue
        if net.version != 4:
            continue
        if net.prefixlen < 24:
            for sub in net.subnets(new_prefix=24):
                v4_24_set.add(str(sub))
        else:
            c_start = int(net.network_address) & 0xFFFFFF00
            v4_24_set.add(str(ipaddress.ip_network(f"{ipaddress.IPv4Address(c_start)}/24", strict=False)))
    return len(v4_24_set) * 256


def cal_v6_pfx48_from_prefixes(prefixes: set) -> int:
    """
    IPv6 前缀集合折算为 /48 数量（区间合并去重）。
    注意：与 common/cal_ip_num.py 口径一致：prefixlen > 48 的碎片前缀会被跳过。
    """
    intervals = []
    SHIFT = 80  # 128-48
    for prefix in prefixes:
        try:
            net = ipaddress.ip_network(str(prefix).strip(), strict=False)
        except Exception:
            continue
        if net.version != 6:
            continue
        if net.prefixlen > 48:
            continue
        start_id = int(net.network_address) >> SHIFT
        end_id = int(net.broadcast_address) >> SHIFT
        intervals.append((start_id, end_id))

    if not intervals:
        return 0
    intervals.sort(key=lambda x: x[0])
    curr_start, curr_end = intervals[0]
    merged = []
    for next_start, next_end in intervals[1:]:
        if next_start <= curr_end + 1:
            curr_end = max(curr_end, next_end)
        else:
            merged.append((curr_start, curr_end))
            curr_start, curr_end = next_start, next_end
    merged.append((curr_start, curr_end))

    total_48 = 0
    for s, e in merged:
        total_48 += (e - s + 1)
    return total_48


def get_as_feature_table_name(asn: str, as_info: dict) -> str:
    """
    根据 ASN 推断其对应的 **AS 特征逻辑表名**（不带月份后缀）。

    规则与 core/BGPFeature.insert_to_db 一致：
    - 大国落入 feature_{country_en}
    - 其他落入 feature_other

    注意：当前 backend 已启用 AS 特征“按月表”机制，实际查询会在 database/feature_asn.py
    内根据时间范围自动路由到 {base}_{YYYYMM}，缺月时回退到 {base}_old。
    因此这里仍然返回 base 表名即可。
    """
    asn = str(asn)
    country_cn = as_info.get(asn, {}).get('as_country_cn', '')
    country_en = BIG_COUNTRY.get(country_cn, "")
    return f"feature_{country_en}" if country_en else FEATURE_OTHER_TABLE


def _peak_outage_time_from_as_outage_table(
    conn,
    as_outage_table: str,
    country_cn: str,
    source: str,
    start_time,
    end_time,
    interval_minutes: int = 3,
):
    """
    复用“国家中断详情页-AS中断时序图”的口径：按 interval_minutes 分桶，
    取并发中断 AS 数最多的时间点作为 t_peak。

    Returns:
        tuple[datetime|None, int]: (t_peak, peak_count)
    """
    if isinstance(start_time, str):
        start_time = datetime.datetime.strptime(start_time, "%Y-%m-%d %H:%M:%S")
    if isinstance(end_time, str):
        end_time = datetime.datetime.strptime(end_time, "%Y-%m-%d %H:%M:%S")

    rows = select_as_outage_by_interval(
        conn=conn,
        source=source,
        start_time=start_time,
        end_time=end_time,
        country=country_cn,
        tables=[as_outage_table],
    )
    if not rows:
        return None, 0

    events = []
    for asn, s_time, e_time in rows:
        if not s_time:
            continue
        events.append((s_time, 1, str(asn)))
        if e_time:
            events.append((e_time, -1, str(asn)))

    if not events:
        return None, 0

    # 同一时刻：start 事件优先于 end（与 pandas mergesort 的稳定排序效果一致）
    events.sort(key=lambda x: (x[0], 0 if x[1] == 1 else 1))

    step = datetime.timedelta(minutes=interval_minutes)
    slots = []
    cur = start_time
    while cur <= end_time:
        slots.append(cur)
        cur += step

    active_event_counts = defaultdict(int)
    active_unique_asns = set()
    event_idx = 0

    t_peak = None
    peak_count = -1

    for slot in slots:
        while event_idx < len(events) and events[event_idx][0] <= slot:
            _, change, asn = events[event_idx]
            if change == 1:
                active_event_counts[asn] += 1
                active_unique_asns.add(asn)
            else:
                active_event_counts[asn] -= 1
                if active_event_counts[asn] <= 0:
                    active_event_counts.pop(asn, None)
                    active_unique_asns.discard(asn)
            event_idx += 1

        cnt = len(active_unique_asns)
        if cnt > peak_count:
            peak_count = cnt
            t_peak = slot

    return t_peak, max(0, peak_count)


def _peak_outage_time_from_prefix_outage_table(
    conn,
    prefix_outage_table: str,
    asn: str,
    source: str,
    start_time,
    end_time,
    interval_minutes: int = 3,
):
    """
    对单个 ASN：在 [start_time, end_time] 内按 interval_minutes 分桶，
    取“并发回撤前缀数（唯一 prefix）”最多的时刻作为 t_peak。

    Returns:
        tuple[datetime|None, int]: (t_peak, peak_prefix_count)
    """
    if isinstance(start_time, str):
        start_time = datetime.datetime.strptime(start_time, "%Y-%m-%d %H:%M:%S")
    if isinstance(end_time, str):
        end_time = datetime.datetime.strptime(end_time, "%Y-%m-%d %H:%M:%S")
    if end_time < start_time:
        return None, 0

    rows = select_prefix_outage_by_interval(
        conn=conn,
        source=source,
        start_time=start_time,
        end_time=end_time,
        country=None,
        asn=str(asn),
        tables=[prefix_outage_table],
    )
    if not rows:
        return None, 0

    events = []
    for prefix, s_time, e_time in rows:
        if not prefix or not s_time:
            continue
        events.append((s_time, 1, str(prefix)))
        if e_time:
            events.append((e_time, -1, str(prefix)))

    if not events:
        return None, 0

    # 同一时刻：start 事件优先于 end
    events.sort(key=lambda x: (x[0], 0 if x[1] == 1 else 1))

    step = datetime.timedelta(minutes=interval_minutes)
    slots = []
    cur = start_time
    while cur <= end_time:
        slots.append(cur)
        cur += step

    active_event_counts = defaultdict(int)
    active_unique_prefixes = set()
    event_idx = 0

    t_peak = None
    peak_count = -1

    for slot in slots:
        while event_idx < len(events) and events[event_idx][0] <= slot:
            _, change, prefix = events[event_idx]
            if change == 1:
                active_event_counts[prefix] += 1
                active_unique_prefixes.add(prefix)
            else:
                active_event_counts[prefix] -= 1
                if active_event_counts[prefix] <= 0:
                    active_event_counts.pop(prefix, None)
                    active_unique_prefixes.discard(prefix)
            event_idx += 1

        cnt = len(active_unique_prefixes)
        if cnt > peak_count:
            peak_count = cnt
            t_peak = slot

    return t_peak, max(0, peak_count)


@lru_cache(maxsize=1)
def load_pfx2as_dict() -> dict:
    """加载并缓存 pfx2as_dict（结构: {asn: {prefix: ...}}）。"""
    with open(PFX2AS_DICT_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def build_country_outage_report_data(conn_event, conn_feature, country, start_time, event_id, source, as_info):
    """
    聚合国家中断报告所需的全部数据
    
    :param conn_event: 事件数据库连接 (conn_15)
    :param conn_feature: 特征数据库连接 (conn_11)
    :param country: 国家两字母代码
    :param start_time: 事件开始时间字符串
    :param event_id: 事件ID
    :param source: 数据源
    :param as_info: AS信息字典
    :return: 报告数据字典
    """
    try:
        # 解析时间，构建表名
        if isinstance(start_time, str):
            # 处理URL中的时间格式 (可能是 2026-01-09%2019:45:00 或 2026-01-09 19:45:00)
            start_time_clean = start_time.replace('%20', ' ').replace('+', ' ')
            event_start_dt = datetime.datetime.strptime(start_time_clean, "%Y-%m-%d %H:%M:%S")
        else:
            event_start_dt = start_time
        
        year = event_start_dt.strftime("%Y")
        month = event_start_dt.strftime("%m")
        
        # 表名
        country_outage_table = f"country_outage_{year}{month}"
        as_outage_table = f"as_outage_{year}{month}"
        prefix_outage_table = f"prefix_outage_{year}{month}"
        
        # 1. 获取国家中断事件详情
        event_rows = get_country_outage_de(conn_event, country_outage_table, country, event_id, source)
        if not event_rows:
            return {'status': False, 'msg': '未找到事件记录'}
        
        event_data = event_rows[0]
        country_chinese_name = event_data['country_chinese_name']
        total_as_num = event_data['total_as_num']
        s_time = event_data['s_time']
        e_time = event_data['e_time']
        duration = event_data['duration']
        outage_level = event_data['outage_level']
        max_outage_as_num = event_data['max_outage_as_num']
        outage_level_descr = event_data['outage_level_descr']
        outage_ases = event_data['outage_ases'] or []
        event_info = event_data['event_info']
        
        # 获取国家代码
        country_code = country
        if country_chinese_name in COUNTRY_NAME_TO_CODE:
            country_code = COUNTRY_NAME_TO_CODE[country_chinese_name]
        
        # 2. 时间转换
        beijing_time_str = format_beijing_time(s_time)
        local_time_str, local_tz_name = convert_to_local_time(s_time, country_code)
        
        # 3. 获取基线数据
        baseline = get_country_baseline(
            conn_feature, 
            country_chinese_name, 
            source, 
            s_time,
            'feature_country',
            24  # 取事件前24小时的平均值作为基线
        )
        
        # 4. 获取事件后当前值（最低值）
        current = get_country_current(
            conn_feature,
            country_chinese_name,
            source,
            s_time,
            'feature_country',
            360  # 取事件后360分钟内的最小值
        )
        
        # 5. 计算降幅
        v4_baseline = baseline.get('v4IP_max', 0) or baseline.get('v4IP_num', 0)
        v6_baseline = baseline.get('v6Prefix_max', 0) or baseline.get('v6Prefix_num', 0)
        v4_current = current.get('v4IP_num', 0)
        v6_current = current.get('v6Prefix_num', 0)
        
        v4_drop_percent = calculate_drop_percent(v4_baseline, v4_current)
        v6_drop_percent = calculate_drop_percent(v6_baseline, v6_current)
        
        # 6. 获取时序数据（事件前后各2小时）
        time_series_start = s_time - timedelta(hours=2)
        # 报告场景下，事件结束时间可能偏保守；这里默认向后扩展 1 天用于查询/统计
        # （若 e_time 更晚，则使用更晚的 e_time）
        event_end_time = s_time + REPORT_LOOKAHEAD
        if e_time and e_time > event_end_time:
            event_end_time = e_time
        time_series_end = event_end_time
        time_series = get_country_time_series(
            conn_feature,
            country_chinese_name,
            source,
            time_series_start,
            time_series_end,
            'feature_country'
        )

        # 6.1 确定中断时刻 t_outage：
        # - 优先：从“AS中断时序图”取并发中断AS数最多的时刻
        # - 兜底：从国家 IP 资源时序取 v4IP_num 最小点
        t_outage = s_time
        peak_outage_as_count = 0
        try:
            t_peak, peak_cnt = _peak_outage_time_from_as_outage_table(
                conn=conn_event,
                as_outage_table=as_outage_table,
                country_cn=country_chinese_name,
                source=source,
                start_time=s_time,
                end_time=event_end_time,
                interval_minutes=3,
            )
            if t_peak:
                t_outage = t_peak
                peak_outage_as_count = peak_cnt or 0
        except Exception:
            t_peak = None

        if time_series:
            try:
                # v4IP_num 可能为 0/None；这里优先选择 >0 的最小值
                best = min(
                    time_series,
                    key=lambda x: x.get('v4IP_num', 0) if (x.get('v4IP_num', 0) or 0) > 0 else float('inf')
                )
                if best and best.get('t'):
                    # 仅当未能从 AS 中断时序取到峰值时刻时才使用该兜底口径
                    if not t_peak:
                        t_outage = datetime.datetime.strptime(best['t'], "%Y-%m-%d %H:%M:%S")
            except Exception:
                if not t_peak:
                    t_outage = s_time
        
        # 7. 构建AS回撤状态表数据
        # - total：使用 AS 特征表基线（事件前24小时窗口 max/avg）
        # - outage：按“每个AS自己的中断区间”取其回撤前缀峰值时刻，再折算（IPv4按/24，IPv6按“前缀条目数”）
        as_table = []
        full_outage_count = 0

        outage_ases_str = [str(a) for a in outage_ases]

        # 拉取本次国家事件窗口内的 AS 中断区间（用于每个AS合并多段）
        per_as_intervals = defaultdict(list)
        try:
            as_outage_rows = select_as_outage_by_interval(
                conn=conn_event,
                source=source,
                start_time=s_time,
                end_time=event_end_time,
                country=country_chinese_name,
                tables=[as_outage_table],
            )
            for row_asn, row_s, row_e in as_outage_rows:
                if row_asn and row_s:
                    per_as_intervals[str(row_asn)].append((row_s, row_e))
        except Exception:
            per_as_intervals = defaultdict(list)

        for asn in outage_ases_str:
            intervals = per_as_intervals.get(str(asn), [])
            as_start_time = s_time
            as_end_time = event_end_time
            as_peak_prefix_count = 0

            if intervals:
                as_start_time = min((x[0] for x in intervals if x and x[0]), default=s_time)
                as_end_time = max(((x[1] if x and x[1] else event_end_time) for x in intervals), default=event_end_time)
                try:
                    t_as_peak, peak_cnt = _peak_outage_time_from_prefix_outage_table(
                        conn=conn_event,
                        prefix_outage_table=prefix_outage_table,
                        asn=asn,
                        source=source,
                        start_time=as_start_time,
                        end_time=as_end_time,
                        interval_minutes=3,
                    )
                    if t_as_peak:
                        t_as = t_as_peak
                        as_peak_prefix_count = peak_cnt or 0
                    else:
                        t_as = as_start_time
                except Exception:
                    t_as = as_start_time
            else:
                # 兜底：如果 as_outage 表没有该 ASN 的区间，就按国家 t_outage 计算一次
                t_as = t_outage

            feature_table = get_as_feature_table_name(asn, as_info)
            as_country_cn = (as_info.get(str(asn), {}) or {}).get("as_country_cn", "")
            _debug_log(
                "asn={asn} source={source} as_country_cn={as_country_cn} feature_table={feature_table} "
                "baseline_window=[{b_start}, {b_end})".format(
                    asn=asn,
                    source=source,
                    as_country_cn=as_country_cn,
                    feature_table=feature_table,
                    b_start=(as_start_time - timedelta(hours=24)),
                    b_end=(as_start_time - timedelta(minutes=5)),
                )
            )

            baseline_as = get_as_baseline(
                conn=conn_feature,
                asn=asn,
                source=source,
                event_start_time=as_start_time,
                table_name=feature_table,
                hours_before=24
            )
            _debug_log(f"asn={asn} baseline_as={baseline_as}")

            v4_total = baseline_as.get('v4IP_max', 0) or baseline_as.get('v4IP_num', 0) or 0
            v6_total_prefix = int(baseline_as.get('v6Prefix_max', 0) or baseline_as.get('v6Prefix_num', 0) or 0)
            # 注意：IPv6 地址数非常大，Python int 没问题，但前端 JS 的 Number 会丢精度；
            # 这里对外输出字符串，内部计算仍用 prefix 数口径。
            # v6_total_addr = str(v6_total_prefix * V6_ADDR_PER_48)
            v6_total_addr = str(v6_total_prefix)
            _debug_log(f"asn={asn} v4_total={v4_total} v6_total_prefix={v6_total_prefix}")
            if v4_total == 0 and v6_total_prefix == 0:
                _debug_log(
                    f"asn={asn} baseline_total_is_zero; likely causes: "
                    f"no data in baseline window, wrong feature_table due to as_info/country mapping, "
                    f"or source mismatch."
                )

            outage_prefixes_now = get_outage_prefixes_by_asn_at(
                conn=conn_event,
                table=prefix_outage_table,
                asn=asn,
                t_outage=t_as,
                source=source
            )
            out_v4, out_v6 = _split_prefixes_by_version(outage_prefixes_now)
            v4_outage = cal_v4_addr_from_prefixes(out_v4)
            v6_outage_prefix = int(len(out_v6))
            # v6_outage_addr = str(v6_outage_prefix * V6_ADDR_PER_48)
            v6_outage_addr = str(v6_outage_prefix)

            # 不同口径可能导致轻微超出，做上限保护
            if v4_total > 0:
                v4_outage = min(v4_outage, v4_total)
            if v6_total_prefix > 0:
                v6_outage_prefix = min(v6_outage_prefix, v6_total_prefix)
                # v6_outage_addr = str(v6_outage_prefix * V6_ADDR_PER_48)
                v6_outage_addr = str(v6_outage_prefix)


            v4_ratio = round((v4_outage / v4_total * 100), 2) if v4_total > 0 else 0
            v6_ratio = round((v6_outage_prefix / v6_total_prefix * 100), 2) if v6_total_prefix > 0 else 0

            is_full_v4 = (v4_total > 0 and v4_outage >= v4_total)
            is_full_v6 = (v6_total_prefix > 0 and v6_outage_prefix >= v6_total_prefix)
            is_full = (is_full_v4 and is_full_v6) or (v4_total == 0 and is_full_v6) or (v6_total_prefix == 0 and is_full_v4)
            if is_full:
                full_outage_count += 1

            as_info_data = get_as_info_from_dict(as_info, asn)
            as_table.append({
                'asn': asn,
                'v4_total': v4_total,
                'v6_total': v6_total_addr,
                'v4_outage': v4_outage,
                'v6_outage': v6_outage_addr,
                'v4_ratio': v4_ratio,
                'v6_ratio': v6_ratio,
                'v6_total_prefix': v6_total_prefix,
                'v6_outage_prefix': v6_outage_prefix,
                'asn_outage_start_time': str(as_start_time),
                'asn_outage_end_time': str(as_end_time) if as_end_time else None,
                'asn_outage_time': str(t_as),
                'asn_peak_outage_prefix_count': as_peak_prefix_count,
                'org_name_cn': as_info_data.get('org_name', ''),
                'as_name': as_info_data.get('as_name', ''),
                'country': as_info_data.get('country', '')
            })
        
        # 8. 获取前缀回撤明细（简化：从outage_ases推断）
        prefix_table = []
        prefix_outage_rows = get_country_prefix_outage_details(
            conn_event,
            prefix_outage_table,
            country_code,
            s_time,
            event_end_time,
            source
        )
        
        for row in prefix_outage_rows:
            # 过滤只保留该国家的前缀（通过AS归属判断）
            row_asn = str(row.get('asn', ''))
            if row_asn in [str(a) for a in outage_ases]:
                ip_version, addr_num = _prefix_addr_num(row.get('prefix', ''))
                prefix_table.append({
                    'prefix': row.get('prefix', ''),
                    'asn': row_asn,
                    'withdraw_time': str(row.get('s_time', '')),
                    'recover_time': str(row.get('e_time', '')) if row.get('e_time') else '至今未恢复',
                    'duration': str(row.get('duration', '')) if row.get('duration') else '',
                    'outage_level': row.get('outage_level', ''),
                    'country': row.get('country', ''),
                    'as_name': row.get('as_name', ''),
                    'org_name': row.get('org_name', ''),
                    'ip_version': ip_version,
                    'addr_num': addr_num,
                })
        
        # 9. 获取国家累计分配AS数（从as_info统计）
        allocated_as_count = 0
        visible_as_count = total_as_num
        for asn_key, asn_data in as_info.items():
            if asn_data.get('as_country_cn') == country_chinese_name:
                allocated_as_count += 1
        
        # 10. 组装报告数据
        report_data = {
            'status': True,
            'meta': {
                'country': country,
                'country_code': country_code,
                'country_chinese_name': country_chinese_name,
                'event_id': event_id,
                'source': source,
                'start_time': str(s_time),
                'end_time': str(e_time) if e_time else None,
                'duration': str(duration) if duration else None,
                'outage_level': outage_level,
                'outage_level_descr': outage_level_descr,
                'event_info': event_info,
                'beijing_time_str': beijing_time_str,
                'outage_time': str(t_outage),
                'local_time_str': local_time_str,
                'local_timezone': local_tz_name
            },
            'baseline': {
                'allocated_as_count': allocated_as_count,
                'visible_as_count': visible_as_count,
                'v4IP_num': v4_baseline,
                'v6Prefix_num': v6_baseline,
                'v4Prefix_num': baseline.get('v4Prefix_num', 0)
            },
            'current': {
                'v4IP_num': v4_current,
                'v6Prefix_num': v6_current,
                'visible_as_count': total_as_num - max_outage_as_num
            },
            'impact': {
                'outage_as_count': len(outage_ases),
                'full_outage_as_count': full_outage_count,
                'peak_outage_as_count': peak_outage_as_count,
                'max_outage_as_num': max_outage_as_num,
                'total_as_num': total_as_num,
                'v4_drop_percent': v4_drop_percent,
                'v6_drop_percent': v6_drop_percent,
                'outage_ases': outage_ases
            },
            'time_series': time_series,
            'as_table': as_table,
            'prefix_table': prefix_table
        }
        
        return report_data
        
    except Exception as e:
        traceback.print_exc()
        return {'status': False, 'msg': f'构建报告数据失败: {str(e)}'}


def generate_country_outage_word(report_data, as_info):
    """
    根据报告数据生成Word文档
    
    :param report_data: 报告数据字典
    :param as_info: AS信息字典
    :return: 生成的文件路径
    """
    try:
        doc = Document()
        _apply_doc_base_styles(doc, body_size_pt=18)
        
        meta = report_data.get('meta', {})
        baseline = report_data.get('baseline', {})
        current = report_data.get('current', {})
        impact = report_data.get('impact', {})
        as_table_data = report_data.get('as_table', [])
        prefix_table_data = report_data.get('prefix_table', [])
        time_series = report_data.get('time_series', []) or []
        
        country_name = meta.get('country_chinese_name', '')
        beijing_time = meta.get('beijing_time_str', '')
        local_time = meta.get('local_time_str', '')
        
        # 标题
        title = doc.add_heading(f"关于{country_name}路由中断事件的报告", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 事件概述
        doc.add_paragraph()
        p1 = doc.add_paragraph()
        p1.add_run(f"北京时间{beijing_time}（{country_name}当地时间{local_time}），")
        p1.add_run(f"我中心监测发现{country_name}境内发生大规模路由中断事件。")
        p1.add_run(f"监测显示，受大量IP前缀异常回撤影响，该国互联网外部可达性出现剧烈波动。")
        _set_paragraph_first_line_indent_chars(p1, 2)
        for r in p1.runs:
            _set_run_fonts(r, size=Pt(18))
        
        # 国家资源基线
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.add_run(f"截至{datetime.datetime.now().strftime('%Y年%m月')}，")
        p2.add_run(f"{country_name}累计共分配AS数量{baseline.get('allocated_as_count', 0)}个，")
        p2.add_run(f"其中路由可见AS共计{baseline.get('visible_as_count', 0)}个，")
        p2.add_run(f"全网正常可见的IPv4地址总数为{baseline.get('v4IP_num', 0)}个，")
        p2.add_run(f"IPv6（/48地址块）数量为{baseline.get('v6Prefix_num', 0)}个。")
        _set_paragraph_first_line_indent_chars(p2, 2)
        for r in p2.runs:
            _set_run_fonts(r, size=Pt(18))
        
        # 事件影响汇总
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.add_run(f"经分析，在此次国家级路由中断事件中，")
        p3.add_run(f"共有{impact.get('outage_as_count', 0)}个AS发生路由回撤，")
        p3.add_run(f"其中{impact.get('full_outage_as_count', 0)}个AS处于完全回撤状态。")
        p3.add_run(f"{country_name}当前可见的IPv4地址数降至{current.get('v4IP_num', 0)}个，")
        p3.add_run(f"较正常基准值减少了{impact.get('v4_drop_percent', 0)}%，")
        p3.add_run(f"全网可见的IPv6 /48前缀数量同步减少至{current.get('v6Prefix_num', 0)}个，")
        p3.add_run(f"跌幅达到{impact.get('v6_drop_percent', 0)}%。")
        _set_paragraph_first_line_indent_chars(p3, 2)
        for r in p3.runs:
            _set_run_fonts(r, size=Pt(18))
        
        # 一、监测情况
        doc.add_heading("一、监测情况", level=1)
        
        # （一）可见地址数量变化情况
        doc.add_heading(f"（一）{country_name}可见地址数量变化情况", level=2)
        p4 = doc.add_paragraph()
        p4.add_run(f"监测数据显示，{country_name}全国互联网地址可见性在短时间内出现剧烈波动。")
        p4.add_run(f"IPv6地址数量大幅下降，全网可见/48前缀数由{baseline.get('v6Prefix_num', 0)}个")
        p4.add_run(f"减少至{current.get('v6Prefix_num', 0)}个，")
        v6_visible_ratio = 100 - impact.get('v6_drop_percent', 0)
        p4.add_run(f"只有{v6_visible_ratio:.1f}%地址空间维持可见状态。")
        _set_paragraph_first_line_indent_chars(p4, 2)
        for r in p4.runs:
            _set_run_fonts(r, size=Pt(18))
        
        doc.add_paragraph()
        p5 = doc.add_paragraph()
        v4_baseline_wan = baseline.get('v4IP_num', 0) / 10000
        v4_current_wan = current.get('v4IP_num', 0) / 10000
        v4_visible_ratio = 100 - impact.get('v4_drop_percent', 0)
        p5.add_run(f"与此同时，{country_name}全国可见IPv4地址数量由{v4_baseline_wan:.0f}万")
        p5.add_run(f"减少至{v4_current_wan:.0f}万，约{v4_visible_ratio:.0f}%可见。")
        _set_paragraph_first_line_indent_chars(p5, 2)
        for r in p5.runs:
            _set_run_fonts(r, size=Pt(18))

        # 图1/图2：可见地址数量变化情况
        reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'reports'))
        file_prefix = f"country_outage_{meta.get('country', '')}_{meta.get('event_id', '')}_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
        fig1_path, fig2_path = _build_visible_addr_figures(time_series, reports_dir, file_prefix)
        if fig1_path:
            pic_p1 = doc.add_paragraph()
            pic_p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic_run1 = pic_p1.add_run()
            _set_run_fonts(pic_run1, size=Pt(18))
            pic_run1.add_picture(fig1_path, width=Cm(15))
            cap1 = doc.add_paragraph("图1 IPv6地址前缀数量变化情况")
            cap1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in cap1.runs:
                _set_run_fonts(r, size=Pt(18))
        if fig2_path:
            pic_p2 = doc.add_paragraph()
            pic_p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic_run2 = pic_p2.add_run()
            _set_run_fonts(pic_run2, size=Pt(18))
            pic_run2.add_picture(fig2_path, width=Cm(15))
            cap2 = doc.add_paragraph("图2 IPv4地址数量变化情况")
            cap2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in cap2.runs:
                _set_run_fonts(r, size=Pt(18))
        
        # （二）AS路由回撤情况
        doc.add_heading(f"（二）{country_name}AS路由回撤情况", level=2)
        p6 = doc.add_paragraph()
        p6.add_run(f"北京时间{beijing_time}，{country_name}全国累计共有")
        p6.add_run(f"{impact.get('outage_as_count', 0)}个AS发生路由回撤，统计情况如表1所示。")
        _set_paragraph_first_line_indent_chars(p6, 2)
        for r in p6.runs:
            _set_run_fonts(r, size=Pt(18))
        
        # 表1: AS路由回撤状态表
        if as_table_data:
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            # 表头
            header_cells = table.rows[0].cells
            headers = ['ASN', 'v4_total', 'v6_total', 'v4_outage', 'v6_outage', 'v4_ratio', 'v6_ratio', 'org_name_cn']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                for r in header_cells[i].paragraphs[0].runs:
                    _set_run_fonts(r, size=TABLE_FONT_SIZE)
            
            # 数据行（最多显示20条）
            for row_data in as_table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row_data.get('asn', ''))
                row_cells[1].text = str(row_data.get('v4_total', 0))
                row_cells[2].text = str(row_data.get('v6_total', 0))
                row_cells[3].text = str(row_data.get('v4_outage', 0))
                row_cells[4].text = str(row_data.get('v6_outage', 0))
                row_cells[5].text = f"{row_data.get('v4_ratio', 0)}%"
                row_cells[6].text = f"{row_data.get('v6_ratio', 0)}%"
                row_cells[7].text = str(row_data.get('org_name_cn', ''))

                for cell in row_cells:
                        for para in cell.paragraphs:
                            for r in para.runs:
                                _set_run_fonts(r, size=TABLE_FONT_SIZE)
            
            doc.add_paragraph(f"表1 {country_name}AS路由回撤状态表")
        
        # （三）前缀回撤情况
        doc.add_heading(f"（三）{country_name}前缀回撤情况", level=2)
        p7 = doc.add_paragraph()
        p7.add_run(f"北京时间{beijing_time}，{country_name}前缀回撤情况如表2所示。")
        _set_paragraph_first_line_indent_chars(p7, 2)
        for r in p7.runs:
            _set_run_fonts(r, size=Pt(18))
        
        # 表2: 前缀回撤明细表
        if prefix_table_data:
            doc.add_paragraph()
            headers2 = ['IP段', 'IP版本', '地址数', 'ASN', 'AS名称', '组织', '回撤时间', '结束时间', '等级', '持续时长']
            table2 = doc.add_table(rows=1, cols=len(headers2))
            table2.style = 'Table Grid'
            
            # 表头
            header_cells2 = table2.rows[0].cells
            for i, header in enumerate(headers2):
                header_cells2[i].text = header
                header_cells2[i].paragraphs[0].runs[0].bold = True
                for r in header_cells2[i].paragraphs[0].runs:
                    _set_run_fonts(r, size=TABLE_FONT_SIZE)
            
            # 数据行
            for row_data in prefix_table_data:
                row_cells2 = table2.add_row().cells
                row_cells2[0].text = str(row_data.get('prefix', ''))
                row_cells2[1].text = str(row_data.get('ip_version', ''))
                row_cells2[2].text = str(row_data.get('addr_num', ''))
                row_cells2[3].text = str(row_data.get('asn', ''))
                row_cells2[4].text = str(row_data.get('as_name', ''))
                row_cells2[5].text = str(row_data.get('org_name', ''))
                row_cells2[6].text = str(row_data.get('withdraw_time', ''))
                row_cells2[7].text = str(row_data.get('recover_time', ''))
                row_cells2[8].text = str(row_data.get('outage_level', ''))
                row_cells2[9].text = str(row_data.get('duration', ''))

                for cell in row_cells2:
                        for para in cell.paragraphs:
                            for r in para.runs:
                                _set_run_fonts(r, size=TABLE_FONT_SIZE)
            
            doc.add_paragraph(f"表2 {country_name}前缀回撤情况表")
        
        # 保存文件
        os.makedirs(reports_dir, exist_ok=True)
        file_path = os.path.join(reports_dir, f"{file_prefix}.docx")
        
        doc.save(file_path)
        return file_path
        
    except Exception as e:
        traceback.print_exc()
        return None
