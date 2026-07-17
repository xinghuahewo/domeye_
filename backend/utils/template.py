import os
import sys
import datetime
import matplotlib
import pandas as pd
import ipaddress
from datetime import timedelta
import json
from matplotlib import pyplot as plt
from netaddr import *
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# from config.config import AS_IMPORTANCE_TABLE, FEATURE_TABLE, BOUNDARY_TABLE, CONNECTION_TABLE


from utils.get_as_info import get_as_info
from utils.get_other_info import get_as_importance
from utils.get_prefix_info import get_prefix_domain, get_prefix_domain_auth, get_prefix_domain_num, get_prefix_domain_auth_num, get_prefix_name
from utils.utils import (
    heading_1, heading_2, heading_3, subtitle, text, appendix, generateTbableASInfo, generateTablePrefixInfo,
    generateTableInfo, setTableDetailFormat, draw_leak_graph)



def write_template_1(doc, event_type, d, as_info, prefix_info, domain_info, important_as_dict):
    """
    写模板内容
    """

    level_dict = {
        "high": "高危事件",
        "middle": "中危事件",
        "low": "低危事件"
    }

    # 定义 "事件编号" 的样式（加粗仿宋四号）
    style_title = doc.styles.add_style("EventTitleStyle", WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.size = Pt(14)  # 四号
    style_title.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    style_title.font.name = 'Times New Roman'  # 西文字体
    style_title._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    style_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    # style_title.font.bold = True

    # 定义 para1 的样式（加粗仿宋小四）
    style_content = doc.styles.add_style("EventContentStyle", WD_STYLE_TYPE.PARAGRAPH)
    style_content.font.size = Pt(12)  # 小四
    style_content.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    style_content.font.name = 'Times New Roman'
    # style_content.font.name = "仿宋_GB2312"
    style_content._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    style_content.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    # style_content.font.bold = True

    ### 设置表格内容的字体
    table_content = doc.styles.add_style("TabelContentStyle", WD_STYLE_TYPE.PARAGRAPH)
    table_content.font.size = Pt(10)
    table_content.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    table_content.font.name = 'Times New Roman'
    # style_content.font.name = "仿宋_GB2312"
    table_content._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    table_content.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

    if event_type == 'hijack':
        prefix = d["hijacked_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}AS{}若干前缀遭路由劫持事件报告".format(d['hijacked_as_org'], d['hijacked_as'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        # start_time 2025年05月12日 04时55分38秒
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)

        event_info = ''
        if d['hijacked_as_country'] != None:
            if d['hijacked_as_org'] != None and d['hijacked_as_country'] not in d['hijacked_as_org']:
                event_info += '{}'.format(d['hijacked_as_country'])
        if d['hijacked_as_org'] != None:
            event_info += '{}'.format(d['hijacked_as_org'])
        event_info += '所属AS{}的前缀{}被'.format(d['hijacked_as'], d['hijacked_prefix'])
        if d['hijacker_as_country'] != None:
            if d['hijacker_as_org'] != None and d['hijacker_as_country'] not in d['hijacker_as_org']:
                event_info += '{}'.format(d['hijacker_as_country'])
        if d['hijacker_as_org'] != None:
            event_info += '{}'.format(d['hijacker_as_org'])
        event_info += '所属AS{}劫持。'.format(d['hijacker_as'])

        para1 = f"北京时间{time_text}, {event_info}"
        text(doc, para1)

        # 网站应用服务    DNS权威解析服务器
        domain = get_prefix_domain(prefix_info, d['hijacked_prefix'])
        domain_auth = get_prefix_domain_auth(prefix_info, d['hijacked_prefix'])
        domain_num = get_prefix_domain_num(prefix_info, d['hijacked_prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['hijacked_prefix'])

        hijacker_as_important = get_as_importance(important_as_dict, d['hijacker_as'])
        hijacked_as_important = get_as_importance(important_as_dict, d['hijacked_as'])
        if hijacker_as_important and hijacked_as_important:
            para2 = f"""劫持AS{d['hijacker_as']}和被劫持AS{d['hijacked_as']}均为重要自治系统，"""
        elif hijacker_as_important:
            para2 = f"""劫持AS{d['hijacker_as']}为重要自治系统，"""
        elif hijacked_as_important:
            para2 = f"""被劫持AS{d['hijacked_as']}为重要自治系统，"""
        else:
            para2 = ""

        para2 = para2 + f"被劫持前缀{prefix}"
        if domain_num > 0 and domain_auth_num > 0:
            para2 += f"""包含{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain_auth[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。综合研判为{level}。"""
        elif domain_num > 0:
            para2 += f"""包含{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain_auth[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。综合研判为{level}。"""
        elif domain_auth_num > 0:
            para2 += f"""包含{domain_auth_num}个DNS权威解析服务器，向其发生DNS请求的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关服务，可能造成应用服务中断，导致网站服务不可达。综合研判为{level}。"""
        else:
            para2 += f"不包含网站应用服务或者DNS权威解析服务器。综合研判为{level}。"

        text(doc, para2)


        para3 = get_as_info(as_info, d["hijacked_as"])
        para4 = get_as_info(as_info, d["hijacker_as"])

        text(doc, para3)
        text(doc, para4)


        para5 = f"""针对此次路由劫持事件，建议被劫持方发布更细粒度路由或者协调上游服务提供商过滤劫持报文。此外，该前缀未部署RPKI认证，建议RPKI鉴权。"""
        text(doc, para5)

        generateTbableASInfo(doc, as_info, set([d['hijacked_as'], d["hijacker_as"]]), table_content)

        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)
        
        hijacker_admin = d['hijacker_as_admin']
        if (hijacker_admin):
            para9 = """
        劫持方应急处置联系人："""
            appendix(doc, para9)
            admin_num = len(hijacker_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "劫持方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacker_admin[index]['name'], hijacker_admin[index]['tel'], hijacker_admin[index]['email']
                line_count += 1

            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
    
    elif event_type == 'sub_hijack':
        prefix = d["hijacked_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS若干前缀遭遇子前缀劫持事件报告".format(d['hijacked_as_org'], d['hijacked_as'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)

        event_info = ''
        if d['hijacked_as_country'] != None:
            if d['hijacked_as_org'] != None and d['hijacked_as_country'] not in d['hijacked_as_org']:
                event_info += '{}'.format(d['hijacked_as_country'])
        if d['hijacked_as_org'] != None:
            event_info += '{}'.format(d['hijacked_as_org'])
        event_info += '所属AS{}的前缀{}被'.format(d['hijacked_as'], d['hijacked_prefix'])
        if d['hijacker_as_country'] != None:
            if d['hijacker_as_org'] != None and d['hijacker_as_country'] not in d['hijacker_as_org']:
                event_info += '{}'.format(d['hijacker_as_country'])
        if d['hijacker_as_org'] != None:
            event_info += '{}'.format(d['hijacker_as_org'])
        event_info += '所属AS{}劫持。'.format(d['hijacker_as'])

        para1 = f"北京时间{time_text}, {event_info}"
        text(doc, para1)


        # 网站应用服务    DNS权威解析服务器
        domain = get_prefix_domain(prefix_info, d['hijacked_prefix'])
        domain_auth = get_prefix_domain_auth(prefix_info, d['hijacked_prefix'])
        domain_num = get_prefix_domain_num(prefix_info, d['hijacked_prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['hijacked_prefix'])

        hijacker_as_important = get_as_importance(important_as_dict, d['hijacker_as'])
        hijacked_as_important = get_as_importance(important_as_dict, d['hijacked_as'])
        if hijacker_as_important and hijacked_as_important:
            para2 = f"""劫持AS{d['hijacker_as']}和被劫持AS{d['hijacked_as']}均为重要自治系统，"""
        elif hijacker_as_important:
            para2 = f"""劫持AS{d['hijacker_as']}为重要自治系统，"""
        elif hijacked_as_important:
            para2 = f"""被劫持AS{d['hijacked_as']}为重要自治系统，"""
        else:
            para2 = ""

        para2 = para2 + f"被劫持前缀{prefix}"
        if domain_num > 0 and domain_auth_num > 0:
            para2 += f"""包含{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。综合研判为{level}。"""
        elif domain_num > 0:
            para2 += f"""包含{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。综合研判为{level}。"""
        elif domain_auth_num > 0:
            para2 += f"""包含{domain_auth_num}个DNS权威解析服务器，向其发生DNS请求的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关服务，可能造成应用服务中断，导致网站服务不可达。综合研判为{level}。"""
        else:
            para2 += f"不包含网站应用服务或者DNS权威解析服务器。综合研判为{level}。"

        text(doc, para2)

        para3 = get_as_info(as_info, d["hijacked_as"])
        para4 = get_as_info(as_info, d["hijacker_as"])

        text(doc, para3)
        text(doc, para4)

        version = None
        try:
            version = ipaddress.ip_network(prefix).version
        except:
            version = 4
        prefix_len = d['hijacked_prefix'].split('/')[-1]
        try:
            if version == 4:
                if int(prefix_len) < 24:
                    suggestion = '建议{}对外宣告更细粒度的/24路由前缀'.format(d['hijacked_as_org'])
                else:
                    suggestion = '无法宣告更细粒度路由，建议直接联系涉事方进行协同处置'
            else:
                if int(prefix_len) < 48:
                    suggestion = '建议{}对外宣告更细粒度的/48路由前缀'.format(d['hijacked_as_org'])
                else:
                    suggestion = '无法宣告更细粒度路由，建议直接联系涉事方进行协同处置'
        except:
            suggestion = '无法宣告更细粒度路由，建议直接联系涉事方进行协同处置'
        para7 = """
        上述劫持前缀长度为/{}，{}。""".format(prefix_len, suggestion)
        text(doc, para7)

        generateTbableASInfo(doc, as_info, set([d['hijacked_as'], d["hijacker_as"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

        hijacker_admin = d['hijacker_as_admin']
        if (hijacker_admin):
            para9 = """劫持方应急处置联系人："""
            appendix(doc, para9)
            admin_num = len(hijacker_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "劫持方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacker_admin[index]['name'], hijacker_admin[index]['tel'], hijacker_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
    
    elif event_type == 'prefix_outage':
        prefix = d["prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}若干前缀遭遇路由中断报告".format(d['org_name'], d['asn'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        
        event_info = ''
        if d['country'] != None:
            if d["org_name"] != None and d['country'] not in d['org_name']:
                event_info += '{}'.format(d['country'])
        if d['org_name'] != None:
            event_info += '{}的'.format(d['org_name'])
        if d['asn'] != None:
            event_info += 'AS{}的'.format(d['asn'])
        event_info += '前缀{}发生中断。'.format(d['prefix'])
        para1 = f"北京时间{time_text}, {event_info}"
        text(doc, para1)

        # 网站应用服务    DNS权威解析服务器
        domain = get_prefix_domain(prefix_info, d['prefix'])
        domain_auth = get_prefix_domain_auth(prefix_info, d['prefix'])
        domain_num = get_prefix_domain_num(prefix_info, d['prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['prefix'])

        para2 = f"受影响前缀{prefix}"
        if domain_num > 0 and domain_auth_num > 0:
            para2 += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
        elif domain_num > 0:
            para2 += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
        elif domain_auth_num > 0:
            para2 += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
        else:
            para2 = f"""访问该前缀的网络流量"""
        para2 += f"""由于前缀中断，而无法到达，导致服务不可达。综合研判为{level}。"""

        text(doc, para2)

        para3 = get_as_info(as_info, d["asn"])
        text(doc, para3)

        para4 = """建议AS{}重新宣告路由前缀{}。""".format(d['asn'], d['prefix'])
        text(doc, para4)

        generateTbableASInfo(doc, as_info, set([d['asn']]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

        hijacked_admin = d['as_admin']
        if (hijacked_admin):
            para8 = """中断方应急处置联系人："""
            appendix(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "中断联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

    elif event_type == 'as_outage':
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}遭遇路由中断事件报告".format(d['org_name'], d['asn'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        
        radio = round(d['outage_prefix_num'] / d['as_prefix_num'], 2)
        event_info = ''
        if d['country'] != None:
            if d["org_name"] != None and d['country'] not in d['org_name']:
                event_info += '{}的'.format(d['country'])
        if d['org_name'] != None:
            event_info += '{}的'.format(d['org_name'])
        event_info += '所属AS{}发生AS中断，'.format(d['asn'])
        event_info += '中断前缀数量为{}，中断前缀占比{}。综合研判为{}。'.format(d['as_prefix_num'], radio, level)
        para1 = f"北京时间{time_text}, {event_info}"
        text(doc, para1)


        para3 = get_as_info(as_info, d["asn"])
        text(doc, para3)

        para4 = """建议AS{}重新宣告路由前缀，检查AS路由设备及路由连接关系。""".format(d['asn'])
        text(doc, para4)


        prefix_num = len(d['outage_prefixes'])
        if prefix_num > 5:
            doc.add_page_break()
            p1 = """附：前缀具体信息见下表""".format(time_text, event_info)
            text(doc, p1)
            table1 = doc.add_table(rows=6, cols=6, style="Table Grid")
            table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
            table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "前缀", "名称", "域名数量", "（部分）域名", "归属AS"
            table1.style.font.size = Pt(12)  # 字体大小15磅
            for i in range(1, 6):
                table1.cell(i, 0).text = str(i)
                outage_prefix = d['outage_prefixes'][i-1]
                table1.cell(i, 1).text = outage_prefix
                if prefix_info.get(outage_prefix):
                    table1.cell(i, 2).text = get_prefix_name(prefix_info, outage_prefix)
                    table1.cell(i, 3).text = str(get_prefix_domain_num(prefix_info, outage_prefix))
                    domain = get_prefix_domain(prefix_info, outage_prefix)
                    domain_str = str(domain[0:4]) if len(domain) > 4 else str(domain)
                    table1.cell(i, 4).text = domain_str
                table1.cell(i, 5).text = d['asn']
            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
            # table1.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 字体颜色
            # table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
        else:
            doc.add_page_break()
            p1 = """AS包含的中断前缀：""".format(time_text, event_info)
            text(doc, p1)
            rows = prefix_num + 1
            table1 = doc.add_table(rows=rows, cols=6, style="Table Grid")
            table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
            table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "前缀", "名称", "域名数量", "（部分）域名", "归属AS"
            table1.style.font.size = Pt(12)  # 字体大小15磅
            for i in range(1, rows):
                table1.cell(i, 0).text = str(i)
                outage_prefix = d['outage_prefixes'][i-1]
                table1.cell(i, 1).text = outage_prefix
                if prefix_info.get(outage_prefix):
                    table1.cell(i, 2).text = get_prefix_name(prefix_info, outage_prefix)
                    table1.cell(i, 3).text = str(get_prefix_domain_num(prefix_info, outage_prefix))
                    domain = get_prefix_domain(prefix_info, outage_prefix)
                    domain_str = str(domain[0:4]) if len(domain) > 4 else str(domain)
                    table1.cell(i, 4).text = domain_str
                table1.cell(i, 5).text = d['asn']

            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content



        hijacked_admin = d['as_admin']
        if (hijacked_admin):
            para8 = """
        中断方应急处置联系人："""
            text(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "中断方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
            
    
    elif event_type == 'country_outage':
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}遭遇路由中断事件报告".format(d['country'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')

        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        
        
        radio = round(d['outage_as_num'] / d['total_as_num'], 2)
        event_info = '{}发生国家中断。{}总共有{}个AS，其中中断AS数量为{}，中断AS占比{}'.format( \
                        d['country'], d['country'], d['total_as_num'], d['outage_as_num'], radio)
        para1 = f"北京时间{time_text}, {event_info}。综合研判为{level}。"
        text(doc, para1)

        para2 = """建议{}检查所属的AS路由设备及路由连接关系。""".format(d['country'])
        text(doc, para2)

        
        para3 = "附：AS具体信息见下表"
        generateTbableASInfo(doc, as_info, set(d['outage_ases']), table_content)

    
    elif event_type == 'leak':
        prefix = d["leak_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}发生路由泄漏事件报告".format(d['leak_by_org'], d['leak_by'])
        heading_1(doc, h1)

        subtitle(doc)

        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)

        event_info = ''
        if d['leak_by_country'] != None:
            if d['leak_by_org'] != None and d['leak_by_country'] not in d['leak_by_org']:
                event_info += '{}'.format(d['leak_by_country'])
        if d['leak_by_org'] != None:
            event_info += '{}的'.format(d['leak_by_org'])
        event_info += 'AS{}将来自上游AS的路由{}泄漏给'.format( \
                                                d['leak_by'], d['leak_prefix'])
        if d['leak_to_country'] != None:
            if d['leak_to_org'] != None and d['leak_to_country'] not in d['leak_to_org']:
                event_info += '{}'.format(d['leak_to_country'])
        if d['leak_to_org'] != None:
            event_info += '{}的'.format(d['leak_to_org'])
        event_info += 'AS{}，影响AS{}的可达性。'.format(d['leak_to'], d['prefix_origin_as'])

        para1 = f"北京时间{time_text}, {event_info}"
        text(doc, para1)

        # 网站应用服务    DNS权威解析服务器
        domain = get_prefix_domain(prefix_info, prefix)
        domain_auth = get_prefix_domain_auth(prefix_info, prefix)
        domain_num = get_prefix_domain_num(prefix_info, prefix)
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, prefix)

        para2 = f"受影响前缀"
        if domain_num > 0 and domain_auth_num > 0:
            para2 += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
        elif domain_num > 0:
            para2 += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
        elif domain_auth_num > 0:
            para2 += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
        else:
            para2 = f"""访问该前缀的网络流量"""

        para2 += f"""在AS{d["leak_by"]}过境,由于该AS并不具备过境传输能力，因此到达受害方AS{d["prefix_origin_as"]}的流量会在泄露方AS阻断，造成流量黑洞，导致服务不可达。"""
        text(doc, para2)


        ori_as_info = get_as_info(as_info, d["prefix_origin_as"])
        leak_by_info = get_as_info(as_info, d["leak_by"])
        leak_to_info = get_as_info(as_info, d["leak_to"])
        text(doc, f"受害方{ori_as_info}")
        text(doc, f"泄漏方{leak_by_info}")
        text(doc, f"传播方{leak_to_info}")


        para4 = """建议路由泄漏方AS{}停止并撤回路由，路由传播方AS{}也同时撤回路由，并告知受影响方AS{}路由的受影响情况。""".format(d['leak_by'], d['leak_to'], d['prefix_origin_as'])
        text(doc, para4)

        generateTbableASInfo(doc, as_info, set([d["prefix_origin_as"], d["leak_by"], d["leak_to"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

        leak_by_admin = d['leak_to_admin']
        if (leak_by_admin):
            para8 = """泄漏方应急处置联系人："""
            appendix(doc, para8)
            admin_num = len(leak_by_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "泄漏方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = leak_by_admin[index]['name'], leak_by_admin[index]['tel'], leak_by_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
        
        leak_to_admin = d['leak_to_admin']
        if (leak_to_admin):
            para9 = """传播方应急处置联系人："""
            appendix(doc, para9)
            admin_num = len(leak_to_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "传播方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = leak_to_admin[index]['name'], leak_to_admin[index]['tel'], leak_to_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content


        ori_as_admin = d['ori_as_admin']
        if (ori_as_admin):
            para8 = """受影响方应急处置联系人："""
            text(doc, para8)
            admin_num = len(ori_as_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "受影响方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = ori_as_admin[index]['name'], ori_as_admin[index]['tel'], ori_as_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

def write_template_2(doc, event_type, d, as_info, prefix_info, domain_info, important_as_dict):
    """
    写模板内容
    """

    level_dict = {
        "high": "高危事件",
        "middle": "中危事件",
        "low": "低危事件"
    }

    # 定义 "事件编号" 的样式（加粗仿宋四号）
    style_title = doc.styles.add_style("EventTitleStyle", WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.size = Pt(14)  # 四号
    style_title.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    style_title.font.name = 'Times New Roman'  # 西文字体
    style_title._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    style_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    style_title.font.bold = True
    # style_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 两端对齐

    # 定义 para1 的样式（加粗仿宋小四）
    style_content = doc.styles.add_style("EventContentStyle", WD_STYLE_TYPE.PARAGRAPH)
    style_content.font.size = Pt(12)  # 小四
    style_content.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    style_content.font.name = 'Times New Roman'
    # style_content.font.name = "仿宋_GB2312"
    style_content._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    style_content.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    # style_content.font.bold = True
    # style_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 两端对齐

    ### 设置表格内容的字体
    table_content = doc.styles.add_style("TabelContentStyle", WD_STYLE_TYPE.PARAGRAPH)
    table_content.font.size = Pt(10)
    table_content.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    table_content.font.name = 'Times New Roman'
    # style_content.font.name = "仿宋_GB2312"
    table_content._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    table_content.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色


    if event_type == 'hijack':
        prefix = d["hijacked_prefix"]
        level = level_dict[d["event_level"]]
         # 标题部分：两行标题
        h1 = "关于{}AS{}若干前缀遭路由劫持事件报告".format(d['hijacked_as_org'], d['hijacked_as'])
        # 使用标题样式（您可根据需要调用 heading_1 或 text）
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-Hijack-{}-{}-{}""".format(d['hijacked_prefix'], d['event_id'], s_time_id)

        generateTableInfo(doc, para1, level, style_title, style_content)

        table2 = doc.add_table(rows=5, cols=2, style="Table Grid")
        table2.autofit = False
        table2.allow_autofit = False

        columns = ["监测情况", "影响范围", "被劫持方情况介绍", "劫持方情况介绍", "应急处置建议"]

        # 添加单元格内容
        for i in range(len(columns)):
            col = columns[i]
            if col == "监测情况":
                time = d["start_time"][:-3]
                dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
                utc_time = dt - timedelta(hours=8)
                utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
                time_text = '{}(UTC时间{})'.format(time, utc_time)

                event_info = ''
                if d['hijacked_as_country'] != None:
                    if d['hijacked_as_org'] != None and d['hijacked_as_country'] not in d['hijacked_as_org']:
                        event_info += '{}'.format(d['hijacked_as_country'])
                if d['hijacked_as_org'] != None:
                    event_info += '{}'.format(d['hijacked_as_org'])
                event_info += '所属AS{}的前缀{}被'.format(d['hijacked_as'], d['hijacked_prefix'])
                if d['hijacker_as_country'] != None:
                    if d['hijacker_as_org'] != None and d['hijacker_as_country'] not in d['hijacker_as_org']:
                        event_info += '{}'.format(d['hijacker_as_country'])
                if d['hijacker_as_org'] != None:
                    event_info += '{}'.format(d['hijacker_as_org'])
                event_info += '所属AS{}劫持。'.format(d['hijacker_as'])


                content = f"北京时间{time_text}, {event_info}"

            elif col == "影响范围":
                domain = get_prefix_domain(prefix_info, d['hijacked_prefix'])
                domain_auth = get_prefix_domain_auth(prefix_info, d['hijacked_prefix'])
                domain_num = get_prefix_domain_num(prefix_info, d['hijacked_prefix'])
                domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['hijacked_prefix'])

                hijacker_as_important = get_as_importance(important_as_dict, d['hijacker_as'])
                hijacked_as_important = get_as_importance(important_as_dict, d['hijacked_as'])
                if hijacker_as_important and hijacked_as_important:
                    content = f"""劫持AS{d['hijacker_as']}和被劫持AS{d['hijacked_as']}均为重要自治系统，"""
                elif hijacker_as_important:
                    content = f"""劫持AS{d['hijacker_as']}为重要自治系统，"""
                elif hijacked_as_important:
                    content = f"""被劫持AS{d['hijacked_as']}为重要自治系统，"""
                else:
                    content = ""


                content = content + f"""被劫持前缀{prefix}"""
                if domain_num > 0 and domain_auth_num > 0:
                    content += f"""包含{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
                elif domain_num > 0:
                    content += f"""包含{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
                elif domain_auth_num > 0:
                    content += f"""包含{domain_auth_num}个DNS权威解析服务器，向其发生DNS请求的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关服务，可能造成应用服务中断，导致网站服务不可达。"""
                else:
                    content += f"不包含网站应用服务或者DNS权威解析服务器。"
            elif col == "被劫持方情况介绍":
                content = get_as_info(as_info, d["hijacked_as"])
                
            elif col == "劫持方情况介绍":
                content = get_as_info(as_info, d["hijacker_as"])

            elif col == "应急处置建议":
                version = None
                try:
                    version = ipaddress.ip_network(prefix).version
                except:
                    version = 4
                prefix_len = d['hijacked_prefix'].split('/')[-1]
                try:
                    if version == 4:
                        if int(prefix_len) < 24:
                            suggestion = "建议{}对外宣告更细粒度的/24路由前缀".format(d['hijacked_as_org'])
                        else:
                            suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
                    else:
                        if int(prefix_len) < 48:
                            suggestion = "建议{}对外宣告更细粒度的/48路由前缀".format(d['hijacked_as_org'])
                        else:
                            suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
                except:
                    suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
                content = """上述劫持前缀长度为/{}，{}。""".format(prefix_len, suggestion)
            table2.cell(i, 0).text = col
            table2.cell(i, 1).text = content


        setTableDetailFormat(table2, columns, style_title, style_content, "hijack")
        generateTbableASInfo(doc, as_info, set([d['hijacked_as'], d["hijacker_as"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

    elif event_type == 'sub_hijack':
        prefix = d["hijacked_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}若干前缀遭遇子前缀劫持事件报告".format(d['hijacked_as_org'], d['hijacked_as'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-SubHijack-{}-{}-{}""".format(d['hijacked_prefix'], d['event_id'], s_time_id)
        generateTableInfo(doc, para1, level, style_title, style_content)


        table2 = doc.add_table(rows=5, cols=2, style="Table Grid")
        table2.autofit = False
        table2.allow_autofit = False
        # table2.columns[0].width = Inches(3)
        # table2.columns[1].width = Inches(10)

        columns = ["监测情况", "影响范围", "被劫持方情况介绍", "劫持方情况介绍", "应急处置建议"]

        # 添加单元格内容
        for i in range(len(columns)):
            col = columns[i]
            if col == "监测情况":
                
                time = d["start_time"][:-3]
                dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
                utc_time = dt - timedelta(hours=8)
                utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
                time_text = '{}(UTC时间{})'.format(time, utc_time)

                event_info = ''
                if d['hijacked_as_country'] != None:
                    if d['hijacked_as_org'] != None and d['hijacked_as_country'] not in d['hijacked_as_org']:
                        event_info += '{}'.format(d['hijacked_as_country'])
                if d['hijacked_as_org'] != None:
                    event_info += '{}'.format(d['hijacked_as_org'])
                event_info += '所属AS{}的前缀{}被'.format(d['hijacked_as'], d['hijacked_prefix'])
                if d['hijacker_as_country'] != None:
                    if d['hijacker_as_org'] != None and d['hijacker_as_country'] not in d['hijacker_as_org']:
                        event_info += '{}'.format(d['hijacker_as_country'])
                if d['hijacker_as_org'] != None:
                    event_info += '{}'.format(d['hijacker_as_org'])
                event_info += '所属AS{}劫持。'.format(d['hijacker_as'])

                content = f"北京时间{time_text}, {event_info}"

            elif col == "影响范围":
                domain = get_prefix_domain(prefix_info, d['hijacked_prefix'])
                domain_auth = get_prefix_domain_auth(prefix_info, d['hijacked_prefix'])
                domain_num = get_prefix_domain_num(prefix_info, d['hijacked_prefix'])
                domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['hijacked_prefix'])

                hijacker_as_important = get_as_importance(important_as_dict, d['hijacker_as'])
                hijacked_as_important = get_as_importance(important_as_dict, d['hijacked_as'])
                if hijacker_as_important and hijacked_as_important:
                    content = f"""劫持AS{d['hijacker_as']}和被劫持AS{d['hijacked_as']}均为重要自治系统，"""
                elif hijacker_as_important:
                    content = f"""劫持AS{d['hijacker_as']}为重要自治系统，"""
                elif hijacked_as_important:
                    content = f"""被劫持AS{d['hijacked_as']}为重要自治系统，"""
                else:
                    content = ""


                content = content + f"""被劫持前缀{prefix}"""
                if domain_num > 0 and domain_auth_num > 0:
                    content += f"""包含{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
                elif domain_num > 0:
                    content += f"""包含{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
                elif domain_auth_num > 0:
                    content += f"""包含{domain_auth_num}个DNS权威解析服务器，向其发生DNS请求的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关服务，可能造成应用服务中断，导致网站服务不可达。"""
                else:
                    content += f"不包含网站应用服务或者DNS权威解析服务器。"
            elif col == "被劫持方情况介绍":
                content = get_as_info(as_info, d["hijacked_as"])
                
            elif col == "劫持方情况介绍":
                content = get_as_info(as_info, d["hijacker_as"])

            elif col == "应急处置建议":
                version = None
                try:
                    version = ipaddress.ip_network(prefix).version
                except:
                    version = 4
                prefix_len = d['hijacked_prefix'].split('/')[-1]
                try:
                    if version == 4:
                        if int(prefix_len) < 24:
                            suggestion = "建议{}对外宣告更细粒度的/24路由前缀".format(d['hijacked_as_org'])
                        else:
                            suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
                    else:
                        if int(prefix_len) < 48:
                            suggestion = "建议{}对外宣告更细粒度的/48路由前缀".format(d['hijacked_as_org'])
                        else:
                            suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
                except:
                    suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
                content = """上述劫持前缀长度为/{}，{}。""".format(prefix_len, suggestion)
            table2.cell(i, 0).text = col
            table2.cell(i, 1).text = content


        setTableDetailFormat(table2, columns, style_title, style_content, "sub_hijack")
        generateTbableASInfo(doc, as_info, set([d['hijacked_as'], d["hijacker_as"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

    elif event_type == 'prefix_outage':
        prefix = d["prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}路由中断事件报告".format(d['org_name'], d['asn'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-Hijack-{}-{}-{}""".format(d['prefix'], d['event_id'], s_time_id)
        generateTableInfo(doc, para1, level, style_title, style_content)


        table2 = doc.add_table(rows=4, cols=2, style="Table Grid")
        table2.autofit = False
        table2.allow_autofit = False

        columns = ["监测情况", "影响范围", "路由中断方介绍", "应急处置建议"]

        for i in range(len(columns)):
            col = columns[i]
            if col == "监测情况":
                time = d["start_time"][:-3]
                dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
                utc_time = dt - timedelta(hours=8)
                utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
                time_text = '{}(UTC时间{})'.format(time, utc_time)
                
                event_info = ''
                if d['country'] != None:
                    if d["org_name"] != None and d['country'] not in d['org_name']:
                        event_info += '{}'.format(d['country'])
                if d['org_name'] != None:
                    event_info += '{}的'.format(d['org_name'])
                if d['asn'] != None:
                    event_info += 'AS{}的'.format(d['asn'])
                event_info += '前缀{}发生路由回撤中断。'.format(d['prefix'])

                content = f"北京时间{time_text}, {event_info}"

            elif col == "影响范围":
                domain = get_prefix_domain(prefix_info, prefix)
                domain_auth = get_prefix_domain_auth(prefix_info, prefix)
                domain_num = get_prefix_domain_num(prefix_info, prefix)
                domain_auth_num = get_prefix_domain_auth_num(prefix_info, prefix)

                content = f"受影响前缀"
                if domain_num > 0 and domain_auth_num > 0:
                    content += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
                elif domain_num > 0:
                    content += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
                elif domain_auth_num > 0:
                    content += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
                else:
                    content = f"""访问该前缀的网络流量"""

                content += f"""由于前缀中断，而无法到达，导致服务不可达。"""


            elif col == "路由中断方介绍":
                content = get_as_info(as_info, d["asn"])
                
            elif col == "应急处置建议":
                content = """建议AS{}重新宣告路由前缀{}。""".format(d['asn'], d['prefix'])

            table2.cell(i, 0).text = col
            table2.cell(i, 1).text = content






        setTableDetailFormat(table2, columns, style_title, style_content, "outage")

        generateTbableASInfo(doc, as_info, set([d['asn']]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

        # hijacked_admin = d['as_admin']
        # if (hijacked_admin):
        #     para8 = """
        # 中断方应急处置联系人："""
        #     text(doc, para8)
        #     admin_num = len(hijacked_admin)
        #     table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
        #     line_count = 0
        #     for index in range(admin_num):
        #         table.cell(line_count, 0).text = "中断联系人{}".format(line_count+1 if admin_num>1 else "")
        #         table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
        #         table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
        #         = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
        #         line_count += 1

    elif event_type == 'as_outage':
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}遭遇路由中断事件报告".format(d['org_name'], d['asn'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-Outage-{}-{}-{}-{}""".format(d['asn'], d['event_id'], s_time_id, d['event_level'])
        generateTableInfo(doc, para1, level, style_title, style_content)

        table2 = doc.add_table(rows=4, cols=2, style="Table Grid")
        table2.autofit = False
        table2.allow_autofit = False

        columns = ["监测情况", "影响范围", "路由中断方介绍", "应急处置建议"]

        for i in range(len(columns)):
            col = columns[i]
            if col == "监测情况":
                
                time = d["start_time"][:-3]
                dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
                utc_time = dt - timedelta(hours=8)
                utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
                time_text = '{}(UTC时间{})'.format(time, utc_time)
                
                radio = round(d['outage_prefix_num'] / d['as_prefix_num'], 2)
                event_info = ''
                if d['country'] != None:
                    if d["org_name"] != None and d['country'] not in d['org_name']:
                        event_info += '{}的'.format(d['country'])
                if d['org_name'] != None:
                    event_info += '{}的'.format(d['org_name'])
                event_info += '所属AS{}发生AS中断，'.format(d['asn'])
                event_info += '中断前缀数量为{}，中断前缀占比{}'.format(d['as_prefix_num'], radio)

                content = f"北京时间{time_text}, {event_info}。"

            # TODO
            elif col == "影响范围":
                content = ""
                # domain = get_prefix_domain(prefix_info, prefix)
                # domain_auth = get_prefix_domain_auth(prefix_info, prefix)
                # domain_num = get_prefix_domain_num(prefix_info, prefix)
                # domain_auth_num = get_prefix_domain_auth_num(prefix_info, prefix)

                # content = f"受影响前缀"
                # if domain_num > 0 and domain_auth_num > 0:
                #     content += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
                # elif domain_num > 0:
                #     content += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
                # elif domain_auth_num > 0:
                #     content += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
                # else:
                #     content += f"""不包含网站应用服务和DNS权威解析服务器，访问该前缀的网络流量"""

                # content += f"""由于前缀中断，而无法到达，导致服务不可达。"""


            elif col == "路由中断方介绍":
                content = get_as_info(as_info, d["asn"])
                
            elif col == "应急处置建议":
                content = """建议AS{}重新宣告路由前缀，检查AS路由设备及路由连接关系。""".format(d['asn'])

            table2.cell(i, 0).text = col
            table2.cell(i, 1).text = content
        
        setTableDetailFormat(table2, columns, style_title, style_content, "outage")


        doc.add_page_break()
        prefix_num = len(d['outage_prefixes'])
        if prefix_num > 5:
            p1 = """附：AS包含的部分中断前缀：""".format(time_text, event_info)
            text(doc, p1)
            table1 = doc.add_table(rows=6, cols=6, style="Table Grid")
            table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
            table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "前缀", "名称", "域名数量", "（部分）域名", "归属AS"
            table1.style.font.size = Pt(12)  # 字体大小15磅
            for i in range(1, 6):
                table1.cell(i, 0).text = str(i)
                outage_prefix = d['outage_prefixes'][i-1]
                table1.cell(i, 1).text = outage_prefix
                if prefix_info.get(outage_prefix):
                    table1.cell(i, 2).text = get_prefix_name(prefix_info, outage_prefix)
                    table1.cell(i, 3).text = str(get_prefix_domain_num(prefix_info, outage_prefix))
                    domain = get_prefix_domain(prefix_info, outage_prefix)
                    domain_str = str(domain[0:4]) if len(domain) > 4 else str(domain)
                    table1.cell(i, 4).text = domain_str
                table1.cell(i, 5).text = d['asn']
            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = style_content

        else:
            p1 = """附：AS包含的中断前缀：""".format(time_text, event_info)
            text(doc, p1)
            rows = prefix_num + 1
            table1 = doc.add_table(rows=rows, cols=6, style="Table Grid")
            table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
            table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "前缀", "名称", "域名数量", "（部分）域名", "归属AS"
            table1.style.font.size = Pt(12)  # 字体大小15磅
            for i in range(1, rows):
                table1.cell(i, 0).text = str(i)
                outage_prefix = d['outage_prefixes'][i-1]
                table1.cell(i, 1).text = outage_prefix
                if prefix_info.get(outage_prefix):
                    table1.cell(i, 2).text = get_prefix_name(prefix_info, outage_prefix)
                    table1.cell(i, 3).text = str(get_prefix_domain_num(prefix_info, outage_prefix))
                    domain = get_prefix_domain(prefix_info, outage_prefix)
                    domain_str = str(domain[0:4]) if len(domain) > 4 else str(domain)
                    table1.cell(i, 4).text = domain_str
                table1.cell(i, 5).text = d['asn']
            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = style_content
            

        hijacked_admin = d['as_admin']
        if (hijacked_admin):
            para8 = """中断方应急处置联系人："""
            text(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "中断方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = style_content
    
    elif event_type == 'country_outage':
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}遭遇路由中断事件报告".format(d['country'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-CountryOutage-{}-{}-{}-{}""".format(d['country_code'], d['event_id'], s_time_id, d['event_level'])
        generateTableInfo(doc, para1, level, style_title, style_content)

        table2 = doc.add_table(rows=4, cols=2, style="Table Grid")
        table2.autofit = False
        table2.allow_autofit = False

        columns = ["监测情况", "影响范围", "路由中断方介绍", "应急处置建议"]

        for i in range(len(columns)):
            col = columns[i]
            if col == "监测情况":
                
                time = d["start_time"][:-3]
                dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
                utc_time = dt - timedelta(hours=8)
                utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
                time_text = '{}(UTC时间{})'.format(time, utc_time)
                
                
                radio = round(d['outage_as_num'] / d['total_as_num'], 2)
                event_info = '{}发生国家中断。{}总共有{}个AS，其中中断AS数量为{}，中断AS占比{}。'.format( \
                                d['country'], d['country'], d['total_as_num'], d['outage_as_num'], radio)

                content = f"北京时间{time_text}, {event_info}"

            elif col == "影响范围":
                content = ""
                pass
                # domain = get_prefix_domain(prefix_info, prefix)
                # domain_auth = get_prefix_domain_auth(prefix_info, prefix)
                # domain_num = get_prefix_domain_num(prefix_info, prefix)
                # domain_auth_num = get_prefix_domain_auth_num(prefix_info, prefix)

                # content = f"受影响前缀"
                # if domain_num > 0 and domain_auth_num > 0:
                #     content += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
                # elif domain_num > 0:
                #     content += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
                # elif domain_auth_num > 0:
                #     content += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
                # else:
                #     content += f"""不包含网站应用服务和DNS权威解析服务器，访问该前缀的网络流量"""

                # content += f"""由于前缀中断，而无法到达，导致服务不可达。"""


            elif col == "路由中断方介绍":
                content = """中断方：{}""".format(d['country'])
                # content = get_as_info(as_info, d["asn"])
                
            elif col == "应急处置建议":
                content = """建议{}检查所属的AS路由设备及路由连接关系。""".format(d['country'])

            table2.cell(i, 0).text = col
            table2.cell(i, 1).text = content
        
        setTableDetailFormat(table2, columns, style_title, style_content, "outage")
        generateTbableASInfo(doc, as_info, set(d['outage_ases']), table_content)
    
    elif event_type == 'leak':
        prefix = d["leak_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治系统AS{}发生路由泄漏事件报告".format(d['leak_by_org'], d['leak_by'])
        heading_1(doc, h1)

        subtitle(doc)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-Hijack-{}-{}-{}""".format(d['leak_prefix'], d['event_id'], s_time_id)
        generateTableInfo(doc, para1, level, style_title, style_content)

        table2 = doc.add_table(rows=5, cols=2, style="Table Grid")
        table2.autofit = False
        # table2.columns[0].width = Inches(3)
        # table2.columns[1].width = Inches(10)

        columns = ["监测情况", "影响范围", "路由泄露涉事方介绍", "路由泄露示意图", "应急处置建议"]

        for i in range(len(columns)):
            col = columns[i]
            if col == "监测情况":
                
                time = d["start_time"][:-3]
                dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
                utc_time = dt - timedelta(hours=8)
                utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
                time_text = '{}(UTC时间{})'.format(time, utc_time)
                event_info = ''
                if d['leak_by_country'] != None:
                    if d['leak_by_org'] != None and d['leak_by_country'] not in d['leak_by_org']:
                        event_info += '{}'.format(d['leak_by_country'])
                if d['leak_by_org'] != None:
                    event_info += '{}的'.format(d['leak_by_org'])
                event_info += 'AS{}将来自上游AS的路由{}泄漏给'.format( \
                                                        d['leak_by'], d['leak_prefix'])
                if d['leak_to_country'] != None:
                    if d['leak_to_org'] != None and d['leak_to_country'] not in d['leak_to_org']:
                        event_info += '{}'.format(d['leak_to_country'])
                if d['leak_to_org'] != None:
                    event_info += '{}的'.format(d['leak_to_org'])
                event_info += 'AS{}，影响AS{}的可达性。'.format(d['leak_to'], d['prefix_origin_as'])

                content = f"北京时间{time_text}, {event_info}"
            elif col == "影响范围":
                domain = get_prefix_domain(prefix_info, d['leak_prefix'])
                domain_auth = get_prefix_domain_auth(prefix_info, d['leak_prefix'])
                domain_num = get_prefix_domain_num(prefix_info, d['leak_prefix'])
                domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['leak_prefix'])

                content = f"受影响前缀"
                if domain_num > 0 and domain_auth_num > 0:
                    content += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
                elif domain_num > 0:
                    content += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
                elif domain_auth_num > 0:
                    content += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
                else:
                    content = f"""访问该前缀的网络流量"""

                content += f"""在AS{d["leak_by"]}过境,由于该AS并不具备过境传输能力，因此到达受害方AS{d["prefix_origin_as"]}的流量会在泄露方AS阻断，造成流量黑洞，导致服务不可达。"""

                
            elif col == "路由泄露涉事方介绍":

                ori_as_info = get_as_info(as_info, d["prefix_origin_as"])
                leak_by_info = get_as_info(as_info, d["leak_by"])
                leak_to_info = get_as_info(as_info, d["leak_to"])
                
                content = f"受害方{ori_as_info}\n泄漏方{leak_by_info}\n传播方{leak_to_info}"


            elif col == "路由泄露示意图":
                as_path = d["as_path"]
                leak_by = d["leak_by"]
                leak_to = d["leak_to"]
                # 返回生成图片的路径
                content = draw_leak_graph(as_path, leak_by, leak_to, output_img="./reports/leak_graph.png")        
            
            elif col == "应急处置建议":
                content = """建议路由泄漏方AS{}停止并撤回路由，路由传播方AS{}也同时撤回路由，并告知受影响方AS{}路由的受影响情况。""".format(d['leak_by'], d['leak_to'], d['prefix_origin_as'])


            table2.cell(i, 0).text = col
            if col != "路由泄露示意图":
                table2.cell(i, 1).text = content
            else:
                paragraph = table2.cell(i, 1).add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(content, width=Cm(12), height=Cm(5))

        setTableDetailFormat(table2, columns, style_title, style_content, "leak")

        generateTbableASInfo(doc, as_info, set([d["prefix_origin_as"], d["leak_by"], d["leak_to"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)
        # leak_by_admin = d['leak_to_admin']
        # if (leak_by_admin):
        #     para8 = """
        # 泄漏方应急处置联系人："""
        #     text(doc, para8)
        #     admin_num = len(leak_by_admin)
        #     table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
        #     line_count = 0
        #     for index in range(admin_num):
        #         table.cell(line_count, 0).text = "泄漏方联系人{}".format(line_count+1 if admin_num>1 else "")
        #         table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
        #         table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
        #         = leak_by_admin[index]['name'], leak_by_admin[index]['tel'], leak_by_admin[index]['email']
        #         line_count += 1
        
        # leak_to_admin = d['leak_to_admin']
        # if (leak_to_admin):
        #     para9 = """
        # 传播方应急处置联系人："""
        #     text(doc, para9)
        #     admin_num = len(leak_to_admin)
        #     table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
        #     line_count = 0
        #     for index in range(admin_num):
        #         table.cell(line_count, 0).text = "传播方联系人{}".format(line_count+1 if admin_num>1 else "")
        #         table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
        #         table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
        #         = leak_to_admin[index]['name'], leak_to_admin[index]['tel'], leak_to_admin[index]['email']
        #         line_count += 1
        
        # ori_as_admin = d['ori_as_admin']
        # if (ori_as_admin):
        #     para8 = """
        # 受影响方应急处置联系人："""
        #     text(doc, para8)
        #     admin_num = len(ori_as_admin)
        #     table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
        #     line_count = 0
        #     for index in range(admin_num):
        #         table.cell(line_count, 0).text = "受影响方联系人{}".format(line_count+1 if admin_num>1 else "")
        #         table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
        #         table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
        #         = ori_as_admin[index]['name'], ori_as_admin[index]['tel'], ori_as_admin[index]['email']
        #         line_count += 1

def write_template_3(doc, event_type, d, as_info, prefix_info, domain_info, important_as_dict):
    """
    写模板内容
    """
    level_dict = {
        "high": "高危事件",
        "middle": "中危事件",
        "low": "低危事件"
    }
    # 定义 "事件编号" 的样式（加粗仿宋四号）
    style_title = doc.styles.add_style("EventTitleStyle", WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.size = Pt(14)  # 四号
    style_title.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    style_title.font.name = 'Times New Roman'  # 西文字体
    style_title._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    style_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    style_title.font.bold = True

    # 定义 para1 的样式（加粗仿宋小四）
    style_content = doc.styles.add_style("EventContentStyle", WD_STYLE_TYPE.PARAGRAPH)
    style_content.font.size = Pt(12)  # 小四
    style_content.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    style_content.font.name = 'Times New Roman'
    # style_content.font.name = "仿宋_GB2312"
    style_content._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    style_content.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色
    # style_content.font.bold = True

    ### 设置表格内容的字体
    table_content = doc.styles.add_style("TableContentStyle", WD_STYLE_TYPE.PARAGRAPH)
    table_content.font.size = Pt(10)
    table_content.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进
    table_content.font.name = 'Times New Roman'
    # style_content.font.name = "仿宋_GB2312"
    table_content._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")  # 强制指定东亚字体
    table_content.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色


    if event_type == 'hijack':
        prefix = d["hijacked_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}AS{}若干前缀遭路由劫持事件报告".format(d['hijacked_as_org'], d['hijacked_as'])
        heading_1(doc, h1)

        subtitle(doc)

        h2 = "1. 事件描述"
        heading_2(doc, h2)

        h3 = "1.1 事件编号"
        heading_3(doc, h3)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-Hijack-{}-{}-{}-{}""".format(d['hijacked_prefix'], d['event_id'], s_time_id, d['event_level'])
        text(doc, para1)

        h4 = "1.2 监测情况"
        heading_3(doc, h4)

        para2_1 = f"""劫持方：AS{d['hijacker_as']}，{d['hijacker_as_org']}，{d['hijacker_as_country']}"""
        para2_2 = f"""被劫持方：AS{d['hijacked_as']}，{d['hijacked_as_org']}，{d['hijacked_as_country']}""" 
        text(doc, para2_1)
        text(doc, para2_2)

        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)

        event_info = ''
        if d['hijacked_as_country'] != None:
            if d['hijacked_as_org'] != None and d['hijacked_as_country'] not in d['hijacked_as_org']:
                event_info += '{}'.format(d['hijacked_as_country'])
        if d['hijacked_as_org'] != None:
            event_info += '{}'.format(d['hijacked_as_org'])
        event_info += '所属AS{}的前缀{}被'.format(d['hijacked_as'], d['hijacked_prefix'])
        if d['hijacker_as_country'] != None:
            if d['hijacker_as_org'] != None and d['hijacker_as_country'] not in d['hijacker_as_org']:
                event_info += '{}'.format(d['hijacker_as_country'])
        if d['hijacker_as_org'] != None:
            event_info += '{}'.format(d['hijacker_as_org'])
        event_info += '所属AS{}劫持。'.format(d['hijacker_as'])

        content = f"北京时间{time_text}, {event_info}"

        text(doc, content)

        hijacker_info = get_as_info(as_info, d["hijacker_as"])
        hijacked_info = get_as_info(as_info, d["hijacked_as"])

        text(doc, hijacker_info)
        text(doc, hijacked_info)


        h5 = "1.3 影响范围"
        heading_3(doc, h5)


        domain = get_prefix_domain(prefix_info, d['hijacked_prefix'])
        domain_auth = get_prefix_domain_auth(prefix_info, d['hijacked_prefix'])
        domain_num = get_prefix_domain_num(prefix_info, d['hijacked_prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['hijacked_prefix'])

        hijacker_as_important = get_as_importance(important_as_dict, d['hijacker_as'])
        hijacked_as_important = get_as_importance(important_as_dict, d['hijacked_as'])
        if hijacker_as_important and hijacked_as_important:
            content = f"""劫持AS{d['hijacker_as']}和被劫持AS{d['hijacked_as']}均为重要自治系统，"""
        elif hijacker_as_important:
            content = f"""劫持AS{d['hijacker_as']}为重要自治系统，"""
        elif hijacked_as_important:
            content = f"""被劫持AS{d['hijacked_as']}为重要自治系统，"""
        else:
            content = ""

        para3 = content + f"""被劫持前缀{prefix}"""
        if domain_num > 0 and domain_auth_num > 0:
            para3 += f"""包含{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
        elif domain_num > 0:
            para3 += f"""包含{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
        elif domain_auth_num > 0:
            para3 += f"""包含{domain_auth_num}个DNS权威解析服务器，向其发生DNS请求的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关服务，可能造成应用服务中断，导致网站服务不可达。"""
        else:
            para3 += f"不包含网站应用服务或者DNS权威解析服务器。"

        text(doc, para3)
        
        
        h6 = "2. 处置建议"
        heading_2(doc, h6)

        version = None
        try:
            version = ipaddress.ip_network(prefix).version
        except:
            version = 4
        prefix_len = d['hijacked_prefix'].split('/')[-1]
        try:
            if version == 4:
                if int(prefix_len) < 24:
                    suggestion = "建议{}对外宣告更细粒度的/24路由前缀".format(d['hijacked_as_org'])
                else:
                    suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
            else:
                if int(prefix_len) < 48:
                    suggestion = "建议{}对外宣告更细粒度的/48路由前缀".format(d['hijacked_as_org'])
                else:
                    suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
        except:
            suggestion = "无法宣告更细粒度路由，建议直接联系涉事方进行协同处置"
        para7 = """
        上述劫持前缀长度为/{}，{}。""".format(prefix_len, suggestion)
        text(doc, para7)

        hijacked_admin = d['hijacked_as_admin']
        if (hijacked_admin):
            para8 = """
        被劫持方应急处置联系人："""
            appendix(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "被劫持方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1
            
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
        
        hijacker_admin = d['hijacker_as_admin']
        if (hijacker_admin):
            para9 = """
        劫持方应急处置联系人："""
            appendix(doc, para9)
            admin_num = len(hijacker_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "劫持方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacker_admin[index]['name'], hijacker_admin[index]['tel'], hijacker_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

        generateTbableASInfo(doc, as_info, set([d['hijacked_as'], d["hijacker_as"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)
 
    elif event_type == 'sub_hijack':
        prefix = d["hijacked_prefix"]
        # 标题部分
        h1 = "关于{}自治域AS{}若干前缀遭遇子前缀劫持事件报告".format(d['hijacked_as_org'], d['hijacked_as'])
        heading_1(doc, h1)

        subtitle(doc)

        h2 = "1. 事件描述"
        heading_2(doc, h2)

        h3 = "1.1 事件编号"
        heading_3(doc, h3)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-SubHijack-{}-{}-{}-{}""".format(d['hijacked_prefix'], d['event_id'], s_time_id, d['event_level'])
        text(doc, para1)

        h4 = "1.2 监测情况"
        heading_3(doc, h4)

        para2_1 = f"""劫持方：AS{d['hijacker_as']}，{d['hijacker_as_org']}，{d['hijacker_as_country']}"""
        para2_2 = f"""被劫持方：AS{d['hijacked_as']}，{d['hijacked_as_org']}，{d['hijacked_as_country']}""" 
        text(doc, para2_1)
        text(doc, para2_2)


        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)

        event_info = ''
        if d['hijacked_as_country'] != None:
            if d['hijacked_as_org'] != None and d['hijacked_as_country'] not in d['hijacked_as_org']:
                event_info += '{}'.format(d['hijacked_as_country'])
        if d['hijacked_as_org'] != None:
            event_info += '{}'.format(d['hijacked_as_org'])
        event_info += '所属AS{}的前缀{}被'.format(d['hijacked_as'], d['hijacked_prefix'])
        if d['hijacker_as_country'] != None:
            if d['hijacker_as_org'] != None and d['hijacker_as_country'] not in d['hijacker_as_org']:
                event_info += '{}'.format(d['hijacker_as_country'])
        if d['hijacker_as_org'] != None:
            event_info += '{}'.format(d['hijacker_as_org'])
        event_info += '所属AS{}劫持。'.format(d['hijacker_as'])

        content = f"北京时间{time_text}, {event_info}"

        text(doc, content)


        hijacker_info = get_as_info(as_info, d["hijacker_as"])
        hijacked_info = get_as_info(as_info, d["hijacked_as"])
        # hijacker_info = f"""劫持方{d["hijacker_as_org"]}是一家xxxxxxx公司，该公司拥有xxx个自治系统，AS{d["hijacker_as"]}为其中的一个自治系统，拥有xxx个IPv4前缀，xxx个IPv6前缀，xxx个IPv4对等体，xxx个对等体，AS排名为xxx，类型为xxxx。"""
        # hijacked_info = f"""被劫持方{d["hijacked_as_country"]}的{d["hijacked_as_org"]}是一家xxxx公司，该公司拥有xxx个自治系统，AS{d["hijacked_as"]}为其中的一个自治系统，拥有xxx个IPv4前缀，xxx个IPv6前缀，xxx个IPv4对等体，xxx个对等体，AS排名为xxx，类型为xxxx。"""

        text(doc, hijacker_info)
        text(doc, hijacked_info)


        h7 = "1.3 影响范围"
        heading_3(doc, h7)

        domain = get_prefix_domain(prefix_info, d['hijacked_prefix'])
        domain_auth = get_prefix_domain_auth(prefix_info, d['hijacked_prefix'])
        domain_num = get_prefix_domain_num(prefix_info, d['hijacked_prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['hijacked_prefix'])

        hijacker_as_important = get_as_importance(important_as_dict, d['hijacker_as'])
        hijacked_as_important = get_as_importance(important_as_dict, d['hijacked_as'])
        if hijacker_as_important and hijacked_as_important:
            para2 = f"""劫持AS{d['hijacker_as']}和被劫持AS{d['hijacked_as']}均为重要自治系统，"""
        elif hijacker_as_important:
            para2 = f"""劫持AS{d['hijacker_as']}为重要自治系统，"""
        elif hijacked_as_important:
            para2 = f"""被劫持AS{d['hijacked_as']}为重要自治系统，"""
        else:
            para2 = ""
        
        para3 = para2 + f"""被劫持前缀{prefix}"""
        if domain_num > 0 and domain_auth_num > 0:
            para3 += f"""包含{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
        elif domain_num > 0:
            para3 += f"""包含{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关应用服务程序，可能造成应用服务中断。{domain[0]}等网站的权威服务器域名解析和网站访问可能存在重定向风险，导致网站服务不可达。"""
        elif domain_auth_num > 0:
            para3 += f"""包含{domain_auth_num}个DNS权威解析服务器，向其发生DNS请求的网络流量可能被重定向到AS{d["hijacker_as"]}，由于该自治系统可能不存在相关服务，可能造成应用服务中断，导致网站服务不可达。"""
        else:
            para3 += f"不包含网站应用服务或者DNS权威解析服务器。"

        text(doc, para3)
        

        h9 = "2. 处置建议"
        heading_2(doc, h9)

        version = None
        try:
            version = ipaddress.ip_network(prefix).version
        except:
            version = 4
        prefix_len = d['hijacked_prefix'].split('/')[-1]
        try:
            if version == 4:
                if int(prefix_len) < 24:
                    suggestion = '建议{}对外宣告更细粒度的/24路由前缀'.format(d['hijacked_as_org'])
                else:
                    suggestion = '无法宣告更细粒度路由，建议直接联系涉事方进行协同处置'
            else:
                if int(prefix_len) < 48:
                    suggestion = '建议{}对外宣告更细粒度的/48路由前缀'.format(d['hijacked_as_org'])
                else:
                    suggestion = '无法宣告更细粒度路由，建议直接联系涉事方进行协同处置'
        except:
            suggestion = '无法宣告更细粒度路由，建议直接联系涉事方进行协同处置'
        para7 = """
        上述劫持前缀长度为/{}，{}。""".format(prefix_len, suggestion)
        text(doc, para7)

        hijacked_admin = d['hijacked_as_admin']
        if (hijacked_admin):
            para8 = """
        被劫持方应急处置联系人："""
            text(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "被劫持方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1

            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
        
        hijacker_admin = d['hijacker_as_admin']
        if (hijacker_admin):
            para9 = """
        劫持方应急处置联系人："""
            appendix(doc, para9)
            admin_num = len(hijacker_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "劫持方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacker_admin[index]['name'], hijacker_admin[index]['tel'], hijacker_admin[index]['email']
                line_count += 1

            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

        generateTbableASInfo(doc, as_info, set([d['hijacked_as'], d["hijacker_as"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)
    
    elif event_type == 'prefix_outage':
        prefix = d["prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}的前缀{}遭遇路由中断事件报告".format(d['org_name'], d['asn'], d['prefix'])
        heading_1(doc, h1)

        subtitle(doc)

        h2 = "1. 事件描述"
        heading_2(doc, h2)

        h3 = "1.1 事件编号"
        heading_3(doc, h3)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-PrefixOutage-{}-{}-{}-{}""".format(d['prefix'], d['event_id'], s_time_id, d['event_level'])
        text(doc, para1)

        h4 = "1.2 监测情况"
        heading_3(doc, h4)

        para2 = """中断方：{}，AS{}，{}，{}""".format(d['prefix'], d['asn'], d['org_name'], d['country'])
        text(doc, para2)

       
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        
        event_info = ''
        if d['country'] != None:
            if d["org_name"] != None and d['country'] not in d['org_name']:
                event_info += '{}'.format(d['country'])
        if d['org_name'] != None:
            event_info += '{}的'.format(d['org_name'])
        if d['asn'] != None:
            event_info += 'AS{}的'.format(d['asn'])
        event_info += '前缀{}发生中断。'.format(d['prefix'])
        para1 = f"北京时间{time_text}, {event_info}"
        text(doc, para1)

        outage_as_info = get_as_info(as_info, d["asn"])
        text(doc, outage_as_info)

        h7 = "1.3 影响范围"
        heading_3(doc, h7)

        # 网站应用服务    DNS权威解析服务器
        domain = get_prefix_domain(prefix_info, d['prefix'])
        domain_auth = get_prefix_domain_auth(prefix_info, d['prefix'])
        domain_num = get_prefix_domain_num(prefix_info, d['prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['prefix'])

        para3 = f"受影响前缀{prefix}"
        if domain_num > 0 and domain_auth_num > 0:
            para3 += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
        elif domain_num > 0:
            para3 += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
        elif domain_auth_num > 0:
            para3 += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
        else:
            para3 = f"""访问该前缀的网络流量"""
        para3 += f"""由于前缀中断，而无法到达，导致服务不可达。综合研判为{level}"""

        text(doc, para3)

        h9 = "2. 处置建议"
        heading_2(doc, h9)

        para7 = """
        建议AS{}重新宣告路由前缀{}。""".format(d['asn'], d['prefix'])
        text(doc, para7)

        hijacked_admin = d['as_admin']
        if (hijacked_admin):
            para8 = """
        中断方应急处置联系人："""
            text(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "中断联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

        generateTbableASInfo(doc, as_info, set([d['asn']]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)

    elif event_type == 'as_outage':
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}遭遇路由中断事件报告".format(d['org_name'], d['asn'])
        heading_1(doc, h1)

        subtitle(doc)

        h2 = "1. 事件描述"
        heading_2(doc, h2)

        h3 = "1.1 事件编号"
        heading_3(doc, h3)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """
        CNCERT-BGPMon-Outage-{}-{}-{}-{}
            """.format(d['asn'], d['event_id'], s_time_id, d['event_level'])
        text(doc, para1)

        h4 = "1.2 监测情况"
        heading_3(doc, h4)

        para2 = """中断方：AS{}，{}，{}""".format(d['asn'], d['org_name'], d['country'])
        text(doc, para2)

        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        
        radio = round(d['outage_prefix_num'] / d['as_prefix_num'], 2)
        event_info = ''
        if d['country'] != None:
            if d["org_name"] != None and d['country'] not in d['org_name']:
                event_info += '{}的'.format(d['country'])
        if d['org_name'] != None:
            event_info += '{}的'.format(d['org_name'])
        event_info += '所属AS{}发生AS中断，'.format(d['asn'])
        event_info += '中断前缀数量为{}，中断前缀占比{}。'.format(d['as_prefix_num'], radio)

        para4 = f"北京时间{time_text}, {event_info}"

        outage_as_info = get_as_info(as_info, d["asn"])
        text(doc, outage_as_info)

        h7 = "1.3 影响范围"
        heading_3(doc, h7)

        # para5 = """AS前缀具体信息见下表。""".format(time_text, event_info)
        # text(doc, para5)

        prefix_num = len(d['outage_prefixes'])
        if prefix_num > 5:
            # doc.add_page_break()
            p1 = """AS包含的部分中断前缀：""".format(time_text, event_info)
            text(doc, p1)
            table1 = doc.add_table(rows=6, cols=6, style="Table Grid")
            table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
            table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "前缀", "名称", "域名数量", "（部分）域名", "归属AS"
            table1.style.font.size = Pt(12)  # 字体大小15磅
            for i in range(1, 6):
                table1.cell(i, 0).text = str(i)
                outage_prefix = d['outage_prefixes'][i-1]
                table1.cell(i, 1).text = outage_prefix
                if prefix_info.get(outage_prefix):
                    table1.cell(i, 2).text = get_prefix_name(prefix_info, outage_prefix)
                    table1.cell(i, 3).text = str(get_prefix_domain_num(prefix_info, outage_prefix))
                    domain = get_prefix_domain(prefix_info, outage_prefix)
                    domain_str = str(domain[0:4]) if len(domain) > 4 else str(domain)
                    table1.cell(i, 4).text = domain_str
                table1.cell(i, 5).text = d['asn']
            # table1.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 字体颜色
            # table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

        else:
            # doc.add_page_break()
            p1 = """AS包含的中断前缀：""".format(time_text, event_info)
            text(doc, p1)
            rows = prefix_num + 1
            table1 = doc.add_table(rows=rows, cols=6, style="Table Grid")
            table1.cell(0, 0).text, table1.cell(0, 1).text, table1.cell(0, 2).text, table1.cell(0, 3).text, \
            table1.cell(0, 4).text, table1.cell(0, 5).text = "序号", "前缀", "名称", "域名数量", "（部分）域名", "归属AS"
            table1.style.font.size = Pt(12)  # 字体大小15磅
            for i in range(1, rows):
                table1.cell(i, 0).text = str(i)
                outage_prefix = d['outage_prefixes'][i-1]
                table1.cell(i, 1).text = outage_prefix
                if prefix_info.get(outage_prefix):
                    table1.cell(i, 2).text = get_prefix_name(prefix_info, outage_prefix)
                    table1.cell(i, 3).text = str(get_prefix_domain_num(prefix_info, outage_prefix))
                    domain = get_prefix_domain(prefix_info, outage_prefix)
                    domain_str = str(domain[0:4]) if len(domain) > 4 else str(domain)
                    table1.cell(i, 4).text = domain_str
                table1.cell(i, 5).text = d['asn']
            for row in table1.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
            # table1.style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 字体颜色
            # table1.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐

        h9 = "2. 处置建议"
        heading_2(doc, h9)

        para7 = """建议AS{}重新宣告路由前缀，检查AS路由设备及路由连接关系。""".format(d['asn'])
        text(doc, para7)

        hijacked_admin = d['as_admin']
        if (hijacked_admin):
            para8 = """
        中断方应急处置联系人："""
            text(doc, para8)
            admin_num = len(hijacked_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "中断方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = hijacked_admin[index]['name'], hijacked_admin[index]['tel'], hijacked_admin[index]['email']
                line_count += 1
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
    
    elif event_type == 'country_outage':
        # 标题部分
        h1 = "关于{}遭遇路由中断事件报告".format(d['country'])
        heading_1(doc, h1)

        subtitle(doc)

        h2 = "1. 事件描述"
        heading_2(doc, h2)

        h3 = "1.1 事件编号"
        heading_3(doc, h3)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-CountryOutage-{}-{}-{}-{}""".format(d['country_code'], d['event_id'], s_time_id, d['event_level'])
        text(doc, para1)

        h4 = "1.2 监测情况"
        heading_3(doc, h4)

        para2 = """中断方：{}""".format(d['country'])
        text(doc, para2)

        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        
        
        radio = round(d['outage_as_num'] / d['total_as_num'], 2)
        event_info = '{}发生国家中断。{}总共有{}个AS，其中中断AS数量为{}，中断AS占比{}。'.format( \
                        d['country'], d['country'], d['total_as_num'], d['outage_as_num'], radio)

        para3 = f"北京时间{time_text}, {event_info}"
        text(doc, para3)

        h7 = "1.3 影响范围"
        heading_3(doc, h7)

        para5 = """AS具体信息见附表。""".format(time_text, event_info)
        text(doc, para5)


        h9 = "2. 处置建议"
        heading_2(doc, h9)

        para7 = """建议{}检查所属的AS路由设备及路由连接关系。""".format(d['country'])
        text(doc, para7)
        generateTbableASInfo(doc, as_info, set(d['outage_ases']), table_content)
    
    elif event_type == 'leak':
        prefix = d["leak_prefix"]
        level = level_dict[d["event_level"]]
        # 标题部分
        h1 = "关于{}自治域AS{}发生路由泄漏事件报告".format(d['leak_by_org'], d['leak_by'])
        heading_1(doc, h1)

        subtitle(doc)

        h2 = "1. 事件描述"
        heading_2(doc, h2)

        h3 = "1.1 事件编号"
        heading_3(doc, h3)

        s_time_id = d['start_time'].split('日')[0].replace('年', '').replace('月', '')
        para1 = """CNCERT-BGPMon-Leak-{}-{}-{}-{}""".format(d['leak_prefix'], d['event_id'], s_time_id, d['event_level'])
        text(doc, para1)

        h4 = "1.2 监测情况"
        heading_3(doc, h4)

        para2_1 = f"""泄漏方：AS{d['leak_by']}，{d['leak_by_org']}，{d['leak_by_country']}"""
        para2_2 = f"""传播方：AS{d['leak_to']}，{d['leak_to_org']}，{d['leak_to_country']}"""
        para2_3 = f"""受影响方：AS{d['prefix_origin_as']}，{d['ori_as_org']}，{d['ori_as_country']}"""
        text(doc, para2_1)
        text(doc, para2_2)
        text(doc, para2_3)
        
        time = d["start_time"][:-3]
        dt = datetime.datetime.strptime(time, "%Y年%m月%d日 %H时%M分")
        utc_time = dt - timedelta(hours=8)
        utc_time = utc_time.strftime("%Y年%m月%d日 %H时%M分")
        time_text = '{}(UTC时间{})'.format(time, utc_time)
        event_info = ''
        if d['leak_by_country'] != None:
            if d['leak_by_org'] != None and d['leak_by_country'] not in d['leak_by_org']:
                event_info += '{}'.format(d['leak_by_country'])
        if d['leak_by_org'] != None:
            event_info += '{}的'.format(d['leak_by_org'])
        event_info += 'AS{}将来自上游AS的路由{}泄漏给'.format( \
                                                d['leak_by'], d['leak_prefix'])
        if d['leak_to_country'] != None:
            if d['leak_to_org'] != None and d['leak_to_country'] not in d['leak_to_org']:
                event_info += '{}'.format(d['leak_to_country'])
        if d['leak_to_org'] != None:
            event_info += '{}的'.format(d['leak_to_org'])
        event_info += 'AS{}，影响AS{}的可达性。'.format(d['leak_to'], d['prefix_origin_as'])
        content = f"北京时间{time_text}, {event_info}"

        text(doc, content)


        # ori_as_info = f"""受害方{d["prefix_origin_as"]}是一家xxxxxxx公司，该公司拥有xxx个自治系统，AS{d["prefix_origin_as"]}为其中的一个自治系统，拥有xxx个IPv4前缀，xxx个IPv6前缀，xxx个IPv4对等体，xxx个对等体，AS排名为xxx，类型为xxxx。"""
        # leak_by_info = f"""泄露方{d["leak_by"]}是一家xxxxxxx公司，该公司拥有xxx个自治系统，AS{d["leak_by"]}为其中的一个自治系统，拥有xxx个IPv4前缀，xxx个IPv6前缀，xxx个IPv4对等体，xxx个对等体，AS排名为xxx，类型为xxxx。"""
        # leak_to_info = f"""传播方{d["leak_to"]}的{d["leak_to_org"]}是一家xxxx公司，该公司拥有xxx个自治系统，AS{d["leak_to"]}为其中的一个自治系统，拥有xxx个IPv4前缀，xxx个IPv6前缀，xxx个IPv4对等体，xxx个对等体，AS排名为xxx，类型为xxxx。"""
        
        ori_as_info = get_as_info(as_info, d["prefix_origin_as"])
        leak_by_info = get_as_info(as_info, d["leak_by"])
        leak_to_info = get_as_info(as_info, d["leak_to"])
        
        text(doc, ori_as_info)
        text(doc, leak_by_info)
        text(doc, leak_to_info)
        

        h7 = "1.3 影响范围"
        heading_3(doc, h7)

        domain_num = get_prefix_domain_num(prefix_info, d['leak_prefix'])
        domain_auth_num = get_prefix_domain_auth_num(prefix_info, d['leak_prefix'])
        domain = get_prefix_domain(prefix_info, d["leak_prefix"])
        domain_auth= get_prefix_domain_auth(prefix_info, d['leak_prefix'])

        para2 = f"受影响前缀"
        if domain_num > 0 and domain_auth_num > 0:
            para2 += f"""包含了{domain_num}个网站应用服务和{domain_auth_num}个DNS权威解析服务器，访问{domain[0]}等网站的网络流量"""
        elif domain_num > 0:
            para2 += f"""包含了{domain_num}个网站应用服务，访问{domain[0]}等网站的网络流量"""
        elif domain_auth_num > 0:
            para2 += f"""包含了{domain_auth_num}个DNS权威解析服务器，访问{domain_auth[0]}的网络流量"""
        else:
            para2 = f"""访问该前缀的网络流量"""

        para2 += f"""在AS{d["leak_by"]}过境,由于该AS并不具备过境传输能力，因此到达受害方AS{d["prefix_origin_as"]}的流量会在泄露方AS阻断，造成流量黑洞，导致服务不可达。"""
        text(doc, para2)


        h9 = "2. 处置建议"
        heading_2(doc, h9)


        para7 = """
        建议路由泄漏方AS{}停止并撤回路由，路由传播方AS{}也同时撤回路由，并告知受影响方AS{}路由的受影响情况。""".format(d['leak_by'], d['leak_to'], d['prefix_origin_as'])
        text(doc, para7)

        leak_by_admin = d['leak_to_admin']
        if (leak_by_admin):
            para8 = """
        泄漏方应急处置联系人："""
            text(doc, para8)
            admin_num = len(leak_by_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "泄漏方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = leak_by_admin[index]['name'], leak_by_admin[index]['tel'], leak_by_admin[index]['email']
                line_count += 1

            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content
        
        leak_to_admin = d['leak_to_admin']
        if (leak_to_admin):
            para9 = """
            传播方应急处置联系人："""
            text(doc, para9)
            admin_num = len(leak_to_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "传播方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = leak_to_admin[index]['name'], leak_to_admin[index]['tel'], leak_to_admin[index]['email']
                line_count += 1
        
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

        ori_as_admin = d['ori_as_admin']
        if (ori_as_admin):
            para8 = """
        受影响方应急处置联系人："""
            text(doc, para8)
            admin_num = len(ori_as_admin)
            table = doc.add_table(rows=admin_num, cols=6, style="Table Grid")
            line_count = 0
            for index in range(admin_num):
                table.cell(line_count, 0).text = "受影响方联系人{}".format(line_count+1 if admin_num>1 else "")
                table.cell(line_count, 2).text, table.cell(line_count, 4).text = "联系电话", "电子邮箱"
                table.cell(line_count, 1).text, table.cell(line_count, 3).text, table.cell(line_count, 5).text \
                = ori_as_admin[index]['name'], ori_as_admin[index]['tel'], ori_as_admin[index]['email']
                line_count += 1
            
            # 对于表格的样式进行修改
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].style = table_content

        generateTbableASInfo(doc, as_info, set([d["prefix_origin_as"], d["leak_by"], d["leak_to"]]), table_content)
        generateTablePrefixInfo(doc, domain_info, domain_num, domain_auth_num, domain, domain_auth, table_content)


def __generate_pic(kind, asn, x, y):
    """
    绘制折线图，生成的图片保存在pictures目录中
    :param kind: announcement 或者 withdrawal
    :param asn: AS编号
    :param x: 时间列表
    :param y: 数据列表
    :return:
    """
    if not os.path.exists('pictures'):
        os.makedirs('pictures')
    # 设置字体格式
    matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 用黑体显示中文
    # 设置绘图尺寸
    plt.figure(figsize=(10, 5), dpi=350)
    # 绘制图像
    plt.plot(x, y)
    # 图像优化
    pass
    # 设置标题
    plt.title("AS{} {} Number".format(asn, kind), fontdict={'size': 20})
    # 为两条坐标轴设置名称
    plt.xlabel("Time UTC", fontdict={'size': 18})
    plt.ylabel('{} Num'.format(kind), fontdict={'size': 18})
    # 显示图例
    plt.legend(labels=['{}_num'.format(kind)])
    # 坐标轴旋转
    # plt.xticks(rotation=30)

    plt.xticks(range(0, len(x), 45))
    # 设置底部比例，防止横坐标显示不全
    plt.gcf().subplots_adjust(bottom=0.25)
    plt.savefig("pictures/{}.png".format(kind))

def generate_word(event_type, d):
    print('generate_word函数执行')
    if not os.path.exists('reports'):
        os.makedirs('reports')
    doc = Document()
    if event_type == 'prefix_outage':
        # 增加标题
        doc.add_heading(text='{}(AS{})发生中断'.format(d['outage_prefix'], d['asn']), level=1)

        para1 = """
                在UTC时间{}, 前缀{}({})发生中断，中断持续了{}。
                事件状态为：{}, 研判依据：{}
                研判人：{}, 研判时间：{}。
            """.format(d['start_time'], d['outage_prefix'], d['as_of_prefix'], d['duration'], 
                       d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        # 生成特征图片
        __generate_pic('announcement', d['asn'], d['time_list'], d['announ_list'])
        __generate_pic('withdrawal', d['asn'], d['time_list'], d['withdraw_list'])

        para2 = """
            Announcement数量:
            """
        doc.add_paragraph(para2)
        doc.add_picture('pictures/announcement.png', width=Inches(6.2))

        # 增加分页符
        doc.add_page_break()

        para3 = """
            Withdrawal数量:
            """
        doc.add_paragraph(para3)
        doc.add_picture('pictures/withdrawal.png', width=Inches(6.2))
    elif event_type == 'as_outage':
        # 增加标题
        doc.add_heading(text='AS{}发生中断'.format(d['asn']), level=1)

        para1 = """
                在UTC时间{}, {}发生中断，中断持续了{}。
                该AS共有{}条前缀，其中{}条发生了中断。
                事件状态为：{}, 研判依据：{}
                研判人：{}, 研判时间：{}。
            """.format(d['start_time'], d['outage_as'], d['duration'], d['as_prefix_num'], d['outage_prefix_num'], 
                       d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        # 生成特征图片
        __generate_pic('announcement', d['asn'], d['time_list'], d['announ_list'])
        __generate_pic('withdrawal', d['asn'], d['time_list'], d['withdraw_list'])

        para2 = """
            Announcement数量:
            """
        doc.add_paragraph(para2)
        doc.add_picture('pictures/announcement.png', width=Inches(6.2))

        # 增加分页符
        doc.add_page_break()

        para3 = """
            Withdrawal数量:
            """
        doc.add_paragraph(para3)
        doc.add_picture('pictures/withdrawal.png', width=Inches(6.2))

        para4 = """
            中断的前缀有：
            {}
            """.format(json.dumps(d['outage_prefixes']))
        doc.add_paragraph(para4)
    elif event_type == 'country_outage':
        doc.add_heading(text='{}发生中断'.format(d['outage_country']), level=1)
        para1 = """
            在UTC时间{},{}发生中断，中断持续了{}。
            {}共有{}个AS，其中{}个AS发生了中断。
            事件状态为：{}, 研判依据：{}
            研判人：{}, 研判时间：{}。
            """.format(d['start_time'], d['outage_country'], d['duration'], d['outage_country'], d['total_as_num'], d['outage_as_num'],
                       d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        para2 = """
            中断的AS有:
            {}
        """.format(d['outage_ases'])
        doc.add_paragraph(para2)
    elif event_type == 'hijack':
        doc.add_heading(text='前缀{}发生劫持'.format(d['hijacked_prefix']), level=1)
        para1 = """
            在UTC时间{}, 前缀{}被劫持。被劫持AS：{}，劫持者AS：{}。
            劫持持续了：{}
            事件状态为：{}, 研判依据：{}
            研判人：{}, 研判时间：{}。
            """.format(d['start_time'], d['hijacked_prefix'], d['hijacked_as'], d['hijacker_as'], d['duration'],
                        d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        # 生成特征图片
        __generate_pic('announcement', d['hijacked_asn'], d['time_list'], d['announ_list'])
        __generate_pic('withdrawal', d['hijacked_asn'], d['time_list'], d['withdraw_list'])

        para2 = """
            Announcement数量:
            """
        doc.add_paragraph(para2)
        doc.add_picture('pictures/announcement.png', width=Inches(6.2))

        # 增加分页符
        doc.add_page_break()

        para3 = """
            Withdrawal数量:
            """
        doc.add_paragraph(para3)
        doc.add_picture('pictures/withdrawal.png', width=Inches(6.2))
    elif event_type == 'sub_hijack':
        doc.add_heading(text='前缀{}发生子前缀劫持'.format(d['hijacked_prefix']), level=1)
        para1 = """
            在UTC时间{}, 前缀{}被子前缀{}劫持。被劫持AS：{}，劫持者AS: {}。
            劫持持续了: {}
            事件状态为：{}, 研判依据：{}
            研判人：{}, 研判时间：{}。
        """.format(d['start_time'], d['hijacked_prefix'], d['hijacker_prefix'], d['hijacked_as'], d['hijacker_as'], d['duration'],
                   d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        # 子前缀劫持不显示特征

        # # 生成特征图片
        # __generate_pic('announcement', d['hijacked_asn'], d['time_list'], d['announ_list'])
        # __generate_pic('withdrawal', d['hijacked_asn'], d['time_list'], d['withdraw_list'])
        #
        # para2 = """
        #     Announcement数量:
        #     """
        # doc.add_paragraph(para2)
        # doc.add_picture('pictures/announcement.jpg', width=Inches(6.2))
        #
        # # 增加分页符
        # doc.add_page_break()
        #
        # para3 = """
        #     Withdrawal数量:
        #     """
        # doc.add_paragraph(para3)
        # doc.add_picture('pictures/withdrawal.jpg', width=Inches(6.2))
    elif event_type == 'leak':
        doc.add_heading(text='前缀{}发生泄露'.format(d['leak_prefix']))
        para1 = """
            在UTC时间{}, 前缀{}发生泄露。将路由泄露的AS：{}，泄露给AS：{}
            事件状态为：{}, 研判依据：{}
            研判人：{}, 研判时间：{}。
            """.format(d['event_time'], d['leak_prefix'], d['leak_by_as'], d['leak_to_as'], 
                       d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        para2 = """
            检测出泄露的AS路径：{}
            """.format(d['as_path'])
        doc.add_paragraph(para2)

        # 生成特征图片
        __generate_pic('announcement', d['leak_to'], d['time_list'], d['announ_list'])
        __generate_pic('withdrawal', d['leak_to'], d['time_list'], d['withdraw_list'])

        para3 = """
            Announcement数量:
            """
        doc.add_paragraph(para3)
        doc.add_picture('pictures/announcement.png', width=Inches(6.2))

        # 增加分页符
        doc.add_page_break()

        para4 = """
            Withdrawal数量:
            """
        doc.add_paragraph(para4)
        doc.add_picture('pictures/withdrawal.png', width=Inches(6.2))
    elif event_type == 'boundary_outage':
        # 增加标题
        doc.add_heading(text='边界AS对{}发生中断'.format(d['asn']), level=1)

        para1 = """
                在UTC时间{}, {}发生中断，中断持续了{}。
                该AS共有{}条前缀，其中{}条发生了中断。
                事件状态为：{}, 研判依据：{}
                研判人：{}, 研判时间：{}。
            """.format(d['start_time'], d['outage_as'], d['duration'], d['as_prefix_num'], d['outage_prefix_num'], 
                       d['state'], d['judge_reason'], d['judge_userid'], d['judge_time'])
        doc.add_paragraph(para1)

        # 生成特征图片
        __generate_pic('announcement', d['asn'], d['time_list'], d['announ_list'])
        __generate_pic('withdrawal', d['asn'], d['time_list'], d['withdraw_list'])

        para2 = """
            Announcement数量:
            """
        doc.add_paragraph(para2)
        doc.add_picture('pictures/announcement.png', width=Inches(6.2))

        # 增加分页符
        doc.add_page_break()

        para3 = """
            Withdrawal数量:
            """
        doc.add_paragraph(para3)
        doc.add_picture('pictures/withdrawal.png', width=Inches(6.2))

        para4 = """
            中断的前缀有：
            {}
            """.format(json.dumps(d['outage_prefixes']))
        doc.add_paragraph(para4)

    save_path = os.path.abspath('reports/{}.docx'.format(event_type))
    doc.save(save_path)
    return save_path

def generate_excel(state, data_list):
    print('generate_excel函数执行')
    if not os.path.exists('reports'):
        os.makedirs('reports')
    if state == 'judge':
        df = pd.DataFrame(data_list, columns=['事件状态', '事件类型', '事件等级', '肇事机构', '受害机构', '事件信息', '开始时间', '结束时间'])
    elif state in ['suspected', 'notify', 'misreport']:
        df = pd.DataFrame(data_list, columns=['事件状态', '事件类型', '事件等级', '肇事机构', '受害机构', '事件信息', '开始时间', '结束时间', 
                                                '研判依据', '研判人', '研判时间'])
    elif state == 'notified':
        df = pd.DataFrame(data_list, columns=['事件状态', '事件类型', '事件等级', '肇事机构', '受害机构', '事件信息', '开始时间', '结束时间', 
                                                '研判依据', '研判人', '研判时间', '通报人', '通报时间'])
    elif state == 'abroad':
        df = pd.DataFrame(data_list, columns=['事件类型', '事件等级', '肇事机构', '受害机构', '事件信息', '开始时间', '结束时间'])
    elif state == 'border':
        df = pd.DataFrame(data_list, columns=['起始国AS', '起始国AS国家', '起始国AS名称', '起始国AS所属机构', 
                                                '起始国AS类型', '边界国AS', '边界国AS国家', '边界国AS名称', 
                                                '边界国AS所属机构', '边界国AS类型'])
    elif state == 'connect':
        df = pd.DataFrame(data_list, columns=['起始国', '目的国', '出口AS数量', '入口AS数量', 'AS路径数量', 
                                                '关键路径数量', '机构级关键路径数量', '国家级关键路径数量'])
    save_path = os.path.abspath('reports/{}.xlsx'.format(state))
    df.to_excel(save_path, index=False)
    print(save_path)
    return save_path

def init_as_info(prefix_info_file):
    # global as_info
    df = pd.read_csv(prefix_info_file, keep_default_na=False)
    df.drop_duplicates(subset=['prefix'], keep='first', inplace=True)
    df_new = df.set_index('prefix', drop=True, append=False, inplace=False, verify_integrity=False)
    prefix_info = df_new.to_dict(orient='index')
    domain_num = get_prefix_domain_num(prefix_info, '185.192.56.0/22')
    domain_auth_num = get_prefix_domain_auth_num(prefix_info, '185.192.56.0/22')
    print(domain_num)
    print(domain_auth_num)

# if __name__ == '__main__':
#     from config.database import DATABASE, USER, PASSWORD, HOST, PORT
#     conn_11 = psycopg2.connect(database=DATABASE, user=USER, password=PASSWORD, host=HOST, port=PORT)
#     event_type, d = get_event_info(detail_url="hijack/2025-05-16 16:56:11/103.195.74.0-24/5/r", conn = conn_11)
#     print(event_type)
#     print(d)

#     df = pd.read_csv(IMPORTANT_AS_FILE)
#     df.drop_duplicates(subset=['aut-num'], keep='first', inplace=True)
#     df_new = df.set_index('aut-num', drop=True, append=False, inplace=False, verify_integrity=False)
#     important_as_dict = df_new.to_dict(orient='index')

#     df = pd.read_csv(AS_INFO_FILE, keep_default_na=False, usecols=['asn', 'as_name', 'as_country_cn', 'org_name', 'org_name_cn', 'as_info', 'admin_info', 'tech_info', 'abuse_info',
#                                                                     'type', 'country_rank', 'global_rank'])
#     df.drop_duplicates(subset=['asn'], keep='first', inplace=True)
#     df_new = df.set_index('asn', drop=True, append=False, inplace=False, verify_integrity=False)
#     df_new.index = df_new.index.astype(str)
#     as_info = df_new.to_dict(orient='index')

#     df = pd.read_csv(PREFIX_INFO_FILE, keep_default_na=False, usecols=['prefix', 'name', 'domain_num', 'domain_auth_num', 'domain', 'domain_auth'])
#     df.drop_duplicates(subset=['prefix'], keep='first', inplace=True)
#     df_new = df.set_index('prefix', drop=True, append=False, inplace=False, verify_integrity=False)
#     prefix_info = df_new.to_dict(orient='index')

#     df = pd.read_csv(DOMAIN_INFO_FILE, keep_default_na=False, sep=';', usecols=['url', 'title', 'industry', 'ip', 'ip_prefix', 'auth_ip'])
#     df_cn = pd.read_csv(DOMAIN_CN_INFO_FILE, keep_default_na=False, usecols=['url', 'title', 'industry', 'ip', 'ip_prefix', 'auth_ip'])
#     df = pd.concat([df, df_cn], axis=0, ignore_index=True)
#     df.drop_duplicates(subset=['url'], keep='first', inplace=True)
#     df_new = df.set_index('url', drop=True, append=False, inplace=False, verify_integrity=False)
#     domain_info = df_new.to_dict(orient='index')
#     print("初始化字典")
#     country_info = dict()
#     df = pd.read_excel(COUNTRY_INFO_FILE, keep_default_na=False)
#     for index, row in df.iterrows():
#         chinese_name = row['chinese_short_name']
#         country_info.setdefault(chinese_name, dict())
#         country_info[chinese_name]['two_letter_code'] = row['two_letter_code']
#         country_info[chinese_name]['longitude'] = row['longitude']
#         country_info[chinese_name]['latitude'] = row['latitude']



# # 生成应急处置模板
#     print('generate_template函数执行')
#     if not os.path.exists('reports'):
#         os.makedirs('reports')
#     doc = Document()
#     if index == 1:
#         write_template_1(doc, event_type, d, as_info, prefix_info, domain_info, important_as_dict)
#     elif index == 2:
#         write_template_2(doc, event_type, d, as_info, prefix_info, domain_info, important_as_dict)
#     else:
#         write_template_3(doc, event_type, d, as_info, prefix_info, domain_info, important_as_dict)

#     save_path = os.path.abspath('reports/{}_{}.docx'.format(event_type, index))
#     doc.save(save_path)
